"""Create the fixed DOCX template for the dam patrol daily report."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / "app" / "templates" / "dam_patrol_daily_report_template.docx"


def add_heading(document: Document, text: str, level: int = 1):
    paragraph = document.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return paragraph


def set_cell_text(cell, text: str):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)


def build_template():
    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    styles = document.styles
    styles["Normal"].font.name = "SimSun"
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("坝区安全智能巡查日报")
    title_run.bold = True
    title_run.font.size = Pt(22)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("报告日期：{{ report_date }}    生成时间：{{ generated_at }}")

    add_heading(document, "一、巡查概况", 1)
    summary_table = document.add_table(rows=5, cols=4)
    summary_table.style = "Table Grid"
    summary_items = [
        ("今日事件总数", "{{ stats.total_events }}", "LOW 数量", "{{ stats.low_count }}"),
        ("MEDIUM 数量", "{{ stats.medium_count }}", "HIGH 数量", "{{ stats.high_count }}"),
        ("人员事件数量", "{{ stats.person_event_count }}", "船只/捕鱼事件数量", "{{ stats.boat_fishing_event_count }}"),
        ("自动广播次数", "{{ stats.auto_broadcast_count }}", "人工广播次数", "{{ stats.manual_broadcast_count }}"),
        ("已闭环数量", "{{ stats.closed_count }}", "未闭环数量", "{{ stats.unclosed_count }}"),
    ]
    for row_index, row_items in enumerate(summary_items):
        for col_index, value in enumerate(row_items):
            set_cell_text(summary_table.rows[row_index].cells[col_index], value)

    timing = document.add_paragraph()
    timing.add_run("平均响应时间：").bold = True
    timing.add_run("{{ stats.avg_response_time }}")
    timing.add_run("    平均处置时间：").bold = True
    timing.add_run("{{ stats.avg_disposal_time }}")

    add_heading(document, "二、事件明细", 1)
    detail_table = document.add_table(rows=4, cols=8)
    detail_table.style = "Table Grid"
    headers = ["发生时间", "摄像头", "事件类型", "风险等级", "广播情况", "处置人员", "处置结果", "完成时间"]
    for col_index, header in enumerate(headers):
        set_cell_text(detail_table.rows[0].cells[col_index], header)
    set_cell_text(detail_table.rows[1].cells[0], "{%tr for row in event_rows %}")
    for col_index, variable in enumerate([
        "{{ row.occur_time }}",
        "{{ row.camera_name }}",
        "{{ row.scene_type }}",
        "{{ row.risk_level }}",
        "{{ row.broadcast_status }}",
        "{{ row.operator }}",
        "{{ row.disposal_result }}",
        "{{ row.completed_at }}",
    ]):
        set_cell_text(detail_table.rows[2].cells[col_index], variable)
    set_cell_text(detail_table.rows[3].cells[0], "{%tr endfor %}")

    add_heading(document, "三、重点事件", 1)
    high_table = document.add_table(rows=4, cols=6)
    high_table.style = "Table Grid"
    high_headers = ["发生时间", "摄像头", "事件类型", "风险等级", "处置人员", "处置结果"]
    for col_index, header in enumerate(high_headers):
        set_cell_text(high_table.rows[0].cells[col_index], header)
    set_cell_text(high_table.rows[1].cells[0], "{%tr for row in high_event_rows %}")
    for col_index, variable in enumerate([
        "{{ row.occur_time }}",
        "{{ row.camera_name }}",
        "{{ row.scene_type }}",
        "{{ row.risk_level }}",
        "{{ row.operator }}",
        "{{ row.disposal_result }}",
    ]):
        set_cell_text(high_table.rows[2].cells[col_index], variable)
    set_cell_text(high_table.rows[3].cells[0], "{%tr endfor %}")

    add_heading(document, "四、闭环情况", 1)
    document.add_paragraph("闭环率：{{ stats.closed_rate }}。未闭环事件需继续跟踪处置。")

    add_heading(document, "五、数据来源", 1)
    document.add_paragraph("本报告所有统计数字均来自数据库表：{{ data_sources }}。")
    document.add_paragraph("事件基础信息、处置动作、广播记录按报告日期自动汇总，未调用大模型自由生成事件内容。")

    document.save(TEMPLATE_PATH)
    print(TEMPLATE_PATH)


if __name__ == "__main__":
    build_template()

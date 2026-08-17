"""Editable DOCX renderer for the daily patrol report."""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


BLUE = "0B4F8A"
BLUE_LIGHT = "EAF2F8"
TEXT = "243447"
MUTED = "64748B"
RULE = "D9E2EC"
HIGH = "C62828"
HIGH_BG = "FDECEC"
MEDIUM = "D97706"
MEDIUM_BG = "FFF4E5"
LOW = "356A9A"
LOW_BG = "EAF2F8"
GREEN = "1B7F4B"

SANS = "Noto Sans CJK SC"
SERIF = "Noto Serif CJK SC"

RISK_STYLE = {
    "HIGH": ("高风险", HIGH, HIGH_BG),
    "MEDIUM": ("中风险", MEDIUM, MEDIUM_BG),
    "LOW": ("低风险", LOW, LOW_BG),
}


def _font(run, *, name=SANS, size=9, color=TEXT, bold=False):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:eastAsia"), name)
    language = r_pr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        r_pr.append(language)
    language.set(qn("w:val"), "zh-CN")
    language.set(qn("w:eastAsia"), "zh-CN")
    if r_pr.find(qn("w:noProof")) is None:
        r_pr.append(OxmlElement("w:noProof"))
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def _text(paragraph, text, **kwargs):
    run = paragraph.add_run(str(text))
    _font(run, **kwargs)
    return run


def _paragraph(paragraph, *, align=None, before=0, after=0, line=1.15):
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def _shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _margins(cell, top=90, start=100, bottom=90, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, attrs in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn(f"w:{key}"), str(value))


def _clear_borders(table):
    for row in table.rows:
        for cell in row.cells:
            none = {"val": "nil"}
            _border(cell, top=none, bottom=none, start=none, end=none)


def _crop_picture(paragraph, image_path, *, width, height, crop):
    inline = paragraph.add_run().add_picture(str(image_path), width=width, height=height)
    blip_fill = inline._inline.graphic.graphicData.pic.blipFill
    src_rect = OxmlElement("a:srcRect")
    for key in ("l", "t", "r", "b"):
        src_rect.set(key, str(int(crop.get(key, 0))))
    blip_fill.insert(1, src_rect)


def _asset_path(board_image: Path, filename: str) -> Path | None:
    candidate = board_image.with_name(filename)
    return candidate if candidate.is_file() else None


def _insert_picture(paragraph, image_path: Path, *, width, height=None):
    return paragraph.add_run().add_picture(str(image_path), width=width, height=height)


def _field(paragraph, instruction, result="1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = str(result)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    _font(run, size=8, color=MUTED)


def _bookmark(paragraph, name, bookmark_id="1"):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(start)
    paragraph._p.append(end)


def _page_number_start(section, start=1):
    pg_num = section._sectPr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        section._sectPr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def _page(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(19)
    section.bottom_margin = Mm(17)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(8)


def _header(section, board_image, context):
    report_date = context["report_period_label"]
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Mm(170))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _clear_borders(table)
    left = table.cell(0, 0).paragraphs[0]
    logo_image = _asset_path(board_image, "report_logo.png")
    if logo_image:
        _insert_picture(left, logo_image, width=Mm(33), height=Mm(13.5))
    else:
        _crop_picture(
            left,
            board_image,
            width=Mm(33),
            height=Mm(13.5),
            crop={"l": 1900, "t": 4700, "r": 85600, "b": 86800},
        )
    right = table.cell(0, 1).paragraphs[0]
    _paragraph(right, align=WD_ALIGN_PARAGRAPH.RIGHT, before=0, after=0, line=1)
    _text(right, context.get("report_title", "大藤峡工程空地联动每日处置报告"), size=8, color=MUTED)
    _text(right, f"   {report_date}", size=8, color=BLUE, bold=True)
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        _border(cell, bottom={"val": "single", "sz": "8", "color": BLUE})


def _footer(section, context, numbered=True):
    report_date = context["report_period_label"]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    _paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    _text(p, f"{context.get('period_name', '每日')}处置报告 · {report_date}", size=7.5, color=MUTED)
    if numbered:
        _text(p, "    ·    ", size=7.5, color=RULE)
        _field(p, "PAGE")
        _text(p, "/", size=8, color=MUTED)
        _field(p, "PAGEREF PatrolReportEnd", result="4")


def _section_title(document, number, title, color=BLUE):
    p = document.add_paragraph()
    _paragraph(p, after=5)
    _text(p, f"{number:02d}", name=SERIF, size=17, color=color)
    _text(p, f"  {title}", name=SERIF, size=17, color=color, bold=True)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for key, value in (("val", "single"), ("sz", "14"), ("color", color), ("space", "5")):
        bottom.set(qn(f"w:{key}"), value)
    p_bdr.append(bottom)
    p._p.get_or_add_pPr().append(p_bdr)


def _cover(document, context, board_image):
    section = document.sections[0]
    _page(section)
    section.different_first_page_header_footer = True
    section.top_margin = Mm(13)
    section.bottom_margin = Mm(12)

    logo = document.add_paragraph()
    _paragraph(logo, align=WD_ALIGN_PARAGRAPH.LEFT, after=28)
    logo_image = _asset_path(board_image, "report_logo.png")
    if logo_image:
        _insert_picture(logo, logo_image, width=Mm(52), height=Mm(21.7))
    else:
        _crop_picture(
            logo,
            board_image,
            width=Mm(52),
            height=Mm(21.3),
            crop={"l": 1900, "t": 4700, "r": 85600, "b": 86800},
        )
    title = document.add_paragraph()
    _paragraph(title, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    _text(title, context.get("report_title", "大藤峡工程空地联动每日处置报告"), name=SERIF, size=24, color=BLUE, bold=True)

    rule = document.add_paragraph()
    _paragraph(rule, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    rule_borders = OxmlElement("w:pBdr")
    rule_bottom = OxmlElement("w:bottom")
    for key, value in (("val", "single"), ("sz", "16"), ("color", BLUE), ("space", "2")):
        rule_bottom.set(qn(f"w:{key}"), value)
    rule_borders.append(rule_bottom)
    rule._p.get_or_add_pPr().append(rule_borders)

    photo = document.add_paragraph()
    _paragraph(photo, align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    cover_image = _asset_path(board_image, "report_cover_photo.png")
    if cover_image:
        _insert_picture(photo, cover_image, width=Mm(166), height=Mm(86))
    else:
        _crop_picture(
            photo,
            board_image,
            width=Mm(166),
            height=Mm(86),
            crop={"l": 780, "t": 33860, "r": 74300, "b": 28800},
        )

    info = document.add_table(rows=1, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    _clear_borders(info)
    for cell in info.rows[0].cells:
        _margins(cell, top=130, bottom=130, start=150, end=150)
        _border(cell, top={"val": "single", "sz": "10", "color": BLUE})
    left = info.cell(0, 0).paragraphs[0]
    _paragraph(left, align=WD_ALIGN_PARAGRAPH.LEFT)
    _text(left, "报告日期\n", size=8, color=MUTED)
    _text(left, context["report_period_label"], name=SERIF, size=14, color=BLUE, bold=True)
    right = info.cell(0, 1).paragraphs[0]
    _paragraph(right, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _text(right, "文档编号\n", size=8, color=MUTED)
    _text(right, f"{context.get('document_code_prefix', 'DX-CZBG')}-{context['report_date_compact']}", size=10, color=TEXT, bold=True)


def _contents(document, context, board_image):
    document.add_page_break()
    _header(document.sections[0], board_image, context)
    _footer(document.sections[0], context, numbered=False)
    p = document.add_paragraph()
    _paragraph(p, after=7)
    _text(p, "目录", name=SERIF, size=20, color=BLUE, bold=True)
    intro = document.add_paragraph()
    _paragraph(intro, after=24)
    _text(intro, "CONTENTS", size=8, color=MUTED, bold=True)
    items = [
        ("01", f"{context.get('period_name', '当日')}风险统计", "1"),
        ("02", "高风险事件", "2"),
        ("03", "中风险事件", "3"),
        ("04", "低风险事件", "4"),
    ]
    for number, title, page in items:
        row = document.add_paragraph()
        _paragraph(row, after=17)
        tabs = row.paragraph_format.tab_stops
        tabs.add_tab_stop(Mm(18), WD_TAB_ALIGNMENT.LEFT)
        tabs.add_tab_stop(Mm(168), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        _text(row, number, name=SERIF, size=12, color=BLUE, bold=True)
        _text(row, f"\t{title}\t", name=SERIF, size=11.5, color=TEXT)
        _text(row, page, name=SERIF, size=11.5, color=BLUE)


def _start_body(document, context, board_image):
    section = document.add_section(WD_SECTION.NEW_PAGE)
    _page(section)
    section.different_first_page_header_footer = False
    _page_number_start(section, 1)
    _header(section, board_image, context)
    _footer(section, context)


def _summary(document, context):
    stats = context["stats"]
    _section_title(document, 1, f"{context.get('period_name', '当日')}风险统计")
    overview_title = document.add_paragraph()
    _paragraph(overview_title, before=3, after=5)
    _text(overview_title, "风险概览", name=SERIF, size=11, color=TEXT, bold=True)
    metrics = [
        ("事件总数", stats["total_events"], BLUE, BLUE_LIGHT),
        ("高风险", stats["high_count"], HIGH, HIGH_BG),
        ("中风险", stats["medium_count"], MEDIUM, MEDIUM_BG),
        ("低风险", stats["low_count"], LOW, LOW_BG),
    ]
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, (label, value, color, bg) in enumerate(metrics):
        cell = table.cell(0, index)
        _shade(cell, bg)
        _margins(cell, top=210, bottom=190, start=130, end=130)
        _border(
            cell,
            top={"val": "single", "sz": "8", "color": "FFFFFF"},
            bottom={"val": "single", "sz": "8", "color": "FFFFFF"},
            start={"val": "single", "sz": "8", "color": "FFFFFF"},
            end={"val": "single", "sz": "8", "color": "FFFFFF"},
        )
        p1 = cell.paragraphs[0]
        _paragraph(p1, align=WD_ALIGN_PARAGRAPH.CENTER)
        _text(p1, value, name=SERIF, size=22, color=color, bold=True)
        p2 = cell.add_paragraph()
        _paragraph(p2, align=WD_ALIGN_PARAGRAPH.CENTER)
        _text(p2, label, size=8.5, color=MUTED)

    status_title = document.add_paragraph()
    _paragraph(status_title, before=13, after=5)
    _text(status_title, "处置状态", name=SERIF, size=11, color=TEXT, bold=True)
    status_table = document.add_table(rows=1, cols=2)
    status_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    status_data = [
        ("已闭环", stats["closed_count"], GREEN, "EAF6EF"),
        ("待跟进", stats["open_count"], MEDIUM if stats["open_count"] else MUTED, "F4F6F8"),
    ]
    for index, (label, value, color, bg) in enumerate(status_data):
        cell = status_table.cell(0, index)
        _shade(cell, bg)
        _margins(cell, top=150, bottom=150, start=170, end=170)
        _border(cell, start={"val": "single", "sz": "18", "color": color})
        p = cell.paragraphs[0]
        _paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT)
        _text(p, label, size=9, color=MUTED)
        _text(p, f"    {value} 起", name=SERIF, size=15, color=color, bold=True)

    source_title = document.add_paragraph()
    _paragraph(source_title, before=13, after=5)
    _text(source_title, "事件来源", name=SERIF, size=11, color=TEXT, bold=True)
    source = document.add_table(rows=1, cols=2)
    source.alignment = WD_TABLE_ALIGNMENT.CENTER
    source_data = [
        ("传感器事件", stats["sensor_count"], stats["sensor_rate"]),
        ("视觉检测事件", stats["camera_count"], stats["camera_rate"]),
    ]
    for index, (label, count, rate) in enumerate(source_data):
        cell = source.cell(0, index)
        _shade(cell, "F7F9FB")
        _margins(cell, top=150, bottom=150, start=170, end=170)
        _border(cell, bottom={"val": "single", "sz": "8", "color": RULE})
        p = cell.paragraphs[0]
        _paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT)
        _text(p, label, size=9, color=TEXT, bold=True)
        _text(p, f"\n{count} 起", name=SERIF, size=15, color=BLUE, bold=True)
        _text(p, f"    {rate}", size=8.5, color=MUTED)

    if context.get("period_type") in {"weekly", "monthly"}:
        period_title = document.add_paragraph()
        _paragraph(period_title, before=13, after=5)
        _text(period_title, "周期指标", name=SERIF, size=11, color=TEXT, bold=True)
        period_table = document.add_table(rows=1, cols=3)
        period_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        period_data = [
            ("统计天数", f"{stats['period_days']} 天"),
            ("日均事件", f"{stats['avg_daily_events']} 起"),
            ("闭环率", stats["closed_rate"]),
        ]
        for index, (label, value) in enumerate(period_data):
            cell = period_table.cell(0, index)
            _shade(cell, "F7F9FB")
            _margins(cell, top=150, bottom=150, start=170, end=170)
            _border(cell, bottom={"val": "single", "sz": "8", "color": RULE})
            p = cell.paragraphs[0]
            _paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT)
            _text(p, label, size=9, color=MUTED)
            _text(p, f"\n{value}", name=SERIF, size=14, color=BLUE, bold=True)

    h2 = document.add_paragraph()
    _paragraph(h2, before=13, after=5)
    _text(h2, f"{context.get('period_name', '当日')}结论", name=SERIF, size=11, color=TEXT, bold=True)
    box = document.add_table(rows=1, cols=1)
    cell = box.cell(0, 0)
    _shade(cell, BLUE_LIGHT)
    _margins(cell, top=190, bottom=190, start=190, end=190)
    _border(cell, start={"val": "single", "sz": "22", "color": BLUE})
    _text(cell.paragraphs[0], context["conclusion"], size=9.5, color=TEXT, bold=True)


def _event_block(document, event, risk_key):
    risk_label, color, bg = RISK_STYLE[risk_key]
    table = document.add_table(rows=6, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Mm((27, 58, 27, 58)[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _margins(cell, top=125, bottom=125, start=130, end=130)
            _border(cell, bottom={"val": "single", "sz": "6", "color": RULE})
    head = table.cell(0, 0).merge(table.cell(0, 2))
    _shade(head, bg)
    _margins(head, top=150, bottom=150, start=150, end=150)
    _text(head.paragraphs[0], event["event_name"], name=SERIF, size=12, color=color, bold=True)
    _text(head.paragraphs[0], f"    {event['instance_no']}", size=7.5, color=MUTED)
    badge = table.cell(0, 3)
    _shade(badge, bg)
    _margins(badge, top=150, bottom=150, start=150, end=150)
    badge.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _text(badge.paragraphs[0], risk_label, size=9, color=color, bold=True)

    pairs = [
        ("发生时间", event["occur_time"], "当前状态", event["result_label"]),
        ("来源 / 位置", f"{event['source_label']} · {event['location']}", "完成时间", event["completed_at"]),
    ]
    for row_index, values in enumerate(pairs, 1):
        for col_index in range(4):
            value = values[col_index]
            cell = table.cell(row_index, col_index)
            if col_index in (0, 2) and value:
                _shade(cell, "F7F9FB")
            _text(
                cell.paragraphs[0],
                value or "—",
                size=8.2,
                color=MUTED if col_index in (0, 2) else (GREEN if value == "已闭环" else TEXT),
                bold=col_index in (0, 2) or value == "已闭环",
            )

    details = [
        ("关键观测", event["key_observation"]),
        ("处置情况", event["handling_summary"]),
        ("事件摘要", event["summary"]),
    ]
    for row_index, (label, value) in enumerate(details, 3):
        label_cell = table.cell(row_index, 0)
        value_cell = table.cell(row_index, 1).merge(table.cell(row_index, 3))
        _shade(label_cell, "F7F9FB")
        _text(label_cell.paragraphs[0], label, size=8.2, color=MUTED, bold=True)
        _text(value_cell.paragraphs[0], value or "—", size=8.8, color=TEXT)

    images = event.get("evidence_images") or []
    if images:
        label = document.add_paragraph()
        _paragraph(label, before=8, after=4)
        _text(label, "图像佐证", name=SERIF, size=10, color=TEXT, bold=True)
        for index, evidence in enumerate(images[:2], 1):
            p = document.add_paragraph()
            _paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
            p.add_run().add_picture(io.BytesIO(evidence["content"]), width=Mm(150))
            caption = document.add_paragraph()
            _paragraph(caption, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
            description = evidence.get("description") or f"事件图像 {index}"
            captured_at = evidence.get("captured_at") or ""
            _text(caption, f"{description}{' · ' + captured_at if captured_at else ''}", size=7.5, color=MUTED)
    elif event["source_type"] == "camera":
        p = document.add_paragraph()
        _paragraph(p, before=4, after=5)
        _text(p, "图像佐证：无可用图像", size=8, color=MUTED)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _risk_page(document, context, number, risk_key):
    document.add_page_break()
    label, color, _ = RISK_STYLE[risk_key]
    _section_title(document, number, f"{label}事件", color=color)
    events = context["events_by_risk"][risk_key]
    if not events:
        p = document.add_paragraph()
        _paragraph(p, before=8)
        _text(p, f"{context.get('period_name', '当日')}无{label}事件。", size=10, color=GREEN, bold=True)
        return
    for event in events:
        _event_block(document, event, risk_key)


def render_daily_report_docx(context: dict[str, Any], board_image: Path) -> bytes:
    if not board_image.is_file():
        raise FileNotFoundError(f"report cover asset is missing: {board_image}")
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = SANS
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), SANS)
    normal.font.size = Pt(9)
    document.core_properties.title = f"{context.get('period_name', '每日')}处置报告_EVT_{context['report_date_compact']}"
    document.core_properties.subject = context.get("report_subject", "传感器事件与视觉检测事件每日汇总")
    document.core_properties.author = "box_system"

    _cover(document, context, board_image)
    _contents(document, context, board_image)
    _start_body(document, context, board_image)
    _summary(document, context)
    _risk_page(document, context, 2, "HIGH")
    _risk_page(document, context, 3, "MEDIUM")
    _risk_page(document, context, 4, "LOW")
    _bookmark(document.paragraphs[-1], "PatrolReportEnd")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()

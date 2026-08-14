"""Document text extraction for knowledge indexing."""

from __future__ import annotations

import hashlib
import io
import re
from pathlib import Path
from typing import Any


def normalize_text(text: str) -> str:
    text = text.replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_document_bytes(data: bytes, filename: str) -> str:
    """Extract plain text from common office/document formats.

    Optional parsers are imported lazily so the API still works for txt/md even
    if PDF/XLSX dependencies are not installed in a development environment.
    """
    suffix = Path(filename).suffix.lower().lstrip(".")
    if suffix in {"txt", "md", "csv", "json", "log"}:
        return normalize_text(_decode_text(data))
    if suffix == "docx":
        return normalize_text(_parse_docx(data))
    if suffix == "pdf":
        return normalize_text(_parse_pdf(data))
    if suffix in {"xlsx", "xls"}:
        return normalize_text(_parse_xlsx(data))
    return normalize_text(_decode_text(data))


def parse_document_blocks(data: bytes, filename: str) -> list[dict[str, Any]]:
    """Extract structured text blocks while preserving Word paragraph anchors.

    The plain-text parser is kept for backward compatibility, but knowledge
    indexing needs stable paragraph ids and section paths so reports can cite
    the exact Word source used by the model.
    """
    suffix = Path(filename).suffix.lower().lstrip(".")
    if suffix == "docx":
        return _parse_docx_blocks(data)

    text = parse_document_bytes(data, filename)
    blocks: list[dict[str, Any]] = []
    section_path = ""
    page = None
    for index, paragraph in enumerate(
        [part.strip() for part in re.split(r"\n\s*\n|\r\n\s*\r\n|\n", text) if part.strip()],
        start=1,
    ):
        page_match = re.match(r"^\[第(\d+)页\]\s*(.*)", paragraph, flags=re.S)
        if page_match:
            page = int(page_match.group(1))
            paragraph = page_match.group(2).strip()
            if not paragraph:
                continue
        if _looks_like_section_title(paragraph):
            section_path = paragraph[:255]
        blocks.append(_make_block(index, paragraph, section_path=section_path, source_page=page))
    return blocks


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _parse_docx(data: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - optional dependency guard
        raise RuntimeError("缺少 python-docx，无法解析 DOCX") from exc

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_docx_blocks(data: bytes) -> list[dict[str, Any]]:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - optional dependency guard
        raise RuntimeError("缺少 python-docx，无法解析 DOCX") from exc

    doc = Document(io.BytesIO(data))
    blocks: list[dict[str, Any]] = []
    headings: dict[int, str] = {}
    paragraph_index = 0

    for paragraph in doc.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        paragraph_index += 1
        style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
        heading_level = _heading_level(style_name, text)
        if heading_level:
            headings = {level: value for level, value in headings.items() if level < heading_level}
            headings[heading_level] = text[:255]
        section_path = " > ".join(value for _, value in sorted(headings.items()))
        blocks.append(
            _make_block(
                paragraph_index,
                text,
                section_path=section_path,
                style=style_name,
            )
        )

    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text)]
            if not cells:
                continue
            paragraph_index += 1
            blocks.append(
                _make_block(
                    paragraph_index,
                    " | ".join(cells),
                    section_path=" > ".join(value for _, value in sorted(headings.items())),
                    table_id=f"t-{table_index:04d}",
                    table_row=row_index,
                )
            )
    return blocks


def _make_block(
    index: int,
    text: str,
    *,
    section_path: str = "",
    source_page: int | None = None,
    style: str = "",
    table_id: str | None = None,
    table_row: int | None = None,
) -> dict[str, Any]:
    return {
        "paragraph_id": f"p-{index:04d}",
        "paragraph_index": index,
        "text": text,
        "section_path": section_path[:255],
        "source_page": source_page,
        "style": style,
        "table_id": table_id,
        "table_row": table_row,
        "quote_hash": hashlib.sha256(text.encode("utf-8")).hexdigest(),
    }


def _heading_level(style_name: str, text: str) -> int | None:
    lowered = style_name.lower()
    match = re.search(r"heading\s*(\d+)|标题\s*(\d+)", lowered, flags=re.I)
    if match:
        return int(next(group for group in match.groups() if group))
    if _looks_like_section_title(text):
        depth = text.count(".") + 1 if re.match(r"^\d+(\.\d+){0,4}\s+", text) else 1
        return max(1, min(depth, 6))
    return None


def _looks_like_section_title(text: str) -> bool:
    if len(text) > 90:
        return False
    patterns = [
        r"^第[一二三四五六七八九十百\d]+[章节条]",
        r"^\d+(\.\d+){0,4}\s+",
        r"^[一二三四五六七八九十]+、",
    ]
    return any(re.search(pattern, text) for pattern in patterns)


def _parse_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - optional dependency guard
        raise RuntimeError("缺少 pypdf，无法解析 PDF") from exc

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"\n\n[第{index}页]\n{text}")
    return "\n".join(pages)


def _parse_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except Exception as exc:  # pragma: no cover - optional dependency guard
        raise RuntimeError("缺少 openpyxl，无法解析 Excel") from exc

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"[工作表] {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(value).strip() for value in row if value not in (None, "")]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)

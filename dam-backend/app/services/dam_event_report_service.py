"""Generate event handling reports from DAM workflow LLM results."""

from __future__ import annotations

import datetime as dt
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import urllib.request
import zipfile
from pathlib import Path
from typing import Any, Optional
from urllib.parse import unquote, urlparse, urlunparse
from zoneinfo import ZoneInfo

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Mm, Pt
from docxtpl import DocxTemplate, InlineImage
from loguru import logger
from PIL import Image, ImageDraw, ImageFont, ImageOps
from sqlalchemy.orm import Session, load_only

from app.core.config import BASE_DIR, settings
from app.models.analysis_report import AnalysisReport, AnalysisReportKnowledgeCitation
from app.models.camera import Camera
from app.models.data_source import DataSource
from app.models.event_library import EventLibrary
from app.models.safety_integration import (
    SafetyEventEvidence,
    SafetyEventInstance,
    SafetyEventTimelineLog,
)
from app.services.minio_service import minio_service
from app.services.patrol_report_service import store_generated_document
from app.services.safety_event_runtime_service import safety_event_runtime_service
from app.services.timeline_text import truncate


TEMPLATE_PATH = Path(BASE_DIR) / "app" / "templates" / "dam_event_handling_report_template.docx"
LOCAL_TIMEZONE = ZoneInfo("Asia/Shanghai")
TEMPLATE_FIELDS = {
    "event_name",
    "instance_no",
    "risk_label",
    "result_label",
    "occur_time",
    "completed_at",
    "source_label",
    "location",
    "evidence_count",
    "summary",
    "key_observation",
    "source_summary",
    "handling_source",
    "timeline_count",
    "handling_summary",
    "timeline_summary",
    "evidence_summary",
    "frame_evidence_summary",
    "linkage_evidence_summary",
    "evidence_caption",
    "conclusion",
    "report_time",
    "instance_no_prefix",
    "instance_no_suffix",
    "occur_time_display",
    "event_duration",
    "emergency_level",
    "confidence_label",
    "trigger_summary",
    "screening_summary",
    "model_route_summary",
    "workflow_nodes_summary",
    "specialized_summary",
    "local_analysis_summary",
    "cloud_analysis_summary",
    "scene_detail",
    "risk_assessment_detail",
    "impact_assessment",
    "response_plan",
    "monitoring_suggestions",
    "recommendations_text",
    "evidence_inventory",
    "analysis_limitations",
    "follow_up_actions",
    "knowledge_sources_summary",
}
RISK_NAMES = {"HIGH": "高风险", "MEDIUM": "中风险", "LOW": "低风险"}
STATUS_NAMES = {
    "PENDING": "待处理",
    "PROCESSING": "处理中",
    "COMPLETED": "已完成",
    "FALSE_ALARM": "误报",
}
SYSTEM_FACT_FIELDS = {
    "report_date",
    "report_date_cn",
    "report_time",
    "event_name",
    "instance_no",
    "instance_no_prefix",
    "instance_no_suffix",
    "risk_label",
    "result_label",
    "occur_time",
    "occur_time_display",
    "completed_at",
    "event_duration",
    "source_label",
    "location",
    "evidence_count",
    "timeline_count",
    "timeline_summary",
    "evidence_caption",
    "evidence_image",
    "handling_summary",
    "conclusion",
    "summary",
    "key_observation",
    "source_summary",
    "screening_summary",
    "model_route_summary",
    "frame_evidence_summary",
    "evidence_summary",
}


class DamEventReportService:
    """Render and archive event reports using cloud LLM results with local fallback."""

    def generate_from_workflow(
        self,
        db: Session,
        *,
        instance: SafetyEventInstance,
        event: EventLibrary,
        workflow_payload: dict[str, Any],
    ) -> Optional[AnalysisReport]:
        selected = self.select_llm_report(workflow_payload)
        if not selected:
            logger.info("DAM事件报告跳过：未找到可用的大模型分析结果 instance={}", instance.instance_no)
            return None

        self.sync_instance_risk_from_report(instance, selected)
        db.flush()
        context = self.build_context(db, instance, event, workflow_payload, selected)
        docx_bytes = self.render_docx(context)
        # 报告文件名：以 ECA 触发的事件名命名，如「洪水灾害告警处置报告」
        event_name = getattr(event, "event_name", None) or instance.summary or "安全事件"
        filename = f"{self.safe_filename_part(event_name)}处置报告_{instance.instance_no}.docx"
        document_id = f"dam_event_report_{instance.instance_no}"
        document = store_generated_document(
            user_id=settings.PATROL_REPORT_USER_ID,
            user_name=settings.PATROL_REPORT_USER_NAME,
            document_id=document_id,
            filename=filename,
            content=docx_bytes,
            report_date=context["report_date"],
        )

        report = self.upsert_analysis_report(
            db,
            instance=instance,
            event=event,
            file_url=document["url"],
            report_date=context["report_date"],
        )
        self.store_knowledge_citations(
            db,
            report=report,
            instance=instance,
            workflow_insight=context.get("workflow_insight") or {},
        )
        safety_event_runtime_service.append_timeline(
            db,
            instance,
            action_key=f"dam-event-report:{instance.instance_no}",
            log_type="REPORT",
            trigger_type="AUTO",
            status="SUCCESS",
            message=truncate(f"{event_name}处置报告已生成：{selected['source_label']}（报告编号 {report.id}）"),
            payload={
                "instance_no": instance.instance_no,
                "analysis_report_id": report.id,
                "report_url": document["url"],
                "document_id": document["document_id"],
                "llm_source": selected["source"],
                "llm_source_label": selected["source_label"],
                "cloud_error": selected.get("cloud_error"),
            },
        )
        return report

    def store_knowledge_citations(
        self,
        db: Session,
        *,
        report: AnalysisReport,
        instance: SafetyEventInstance,
        workflow_insight: dict[str, Any],
    ) -> None:
        try:
            db.query(AnalysisReportKnowledgeCitation).filter(
                AnalysisReportKnowledgeCitation.report_id == report.id
            ).delete()
            sources = (
                workflow_insight.get("report_knowledge_sources")
                or workflow_insight.get("knowledge_sources")
            )
            citations = (
                workflow_insight.get("report_sentence_citations")
                or workflow_insight.get("sentence_citations")
            )
            if isinstance(sources, list) and isinstance(citations, list):
                source_by_evidence = {
                    str(item.get("evidence_id") or f"K{item.get('chunk_id')}"): item
                    for item in sources
                    if isinstance(item, dict) and (item.get("evidence_id") or item.get("chunk_id"))
                }
                for citation in citations:
                    if not isinstance(citation, dict):
                        continue
                    sentence = str(citation.get("sentence") or "").strip()
                    if not sentence:
                        continue
                    for evidence_id in citation.get("evidence_ids") or []:
                        evidence_key = str(evidence_id)
                        source = source_by_evidence.get(evidence_key)
                        if not source:
                            continue
                        db.add(AnalysisReportKnowledgeCitation(
                            report_id=report.id,
                            instance_no=instance.instance_no,
                            field_name=str(citation.get("field") or ""),
                            sentence=sentence[:2000],
                            evidence_id=evidence_key,
                            chunk_id=self.integer_or_none(source.get("chunk_id")),
                            document_id=self.integer_or_none(source.get("document_id")),
                            document_title=str(source.get("document_title") or source.get("filename") or "")[:240],
                            section_path=str(source.get("section_path") or "")[:512],
                            clause_id=str(source.get("clause_id") or "")[:128],
                            support_type=str(citation.get("support_type") or "direct")[:32],
                            confidence=str(citation.get("confidence") or "")[:32],
                            citation_json=json.dumps(
                                {"citation": citation, "source": source},
                                ensure_ascii=False,
                                default=str,
                            ),
                        ))
            risk_escalation = workflow_insight.get("risk_escalation")
            risk_hits = risk_escalation.get("knowledge_hits") if isinstance(risk_escalation, dict) else []
            risk_sentence = (
                str(risk_escalation.get("reason") or "").strip()
                if isinstance(risk_escalation, dict)
                else ""
            )
            if isinstance(risk_hits, list) and risk_sentence:
                for hit in risk_hits:
                    if not isinstance(hit, dict):
                        continue
                    evidence_id = str(hit.get("evidence_id") or f"K{hit.get('chunk_id')}" or "").strip()
                    if not evidence_id:
                        continue
                    db.add(AnalysisReportKnowledgeCitation(
                        report_id=report.id,
                        instance_no=instance.instance_no,
                        field_name="risk_escalation",
                        sentence=risk_sentence[:2000],
                        evidence_id=evidence_id,
                        chunk_id=self.integer_or_none(hit.get("chunk_id")),
                        document_id=self.integer_or_none(hit.get("document_id")),
                        document_title=str(hit.get("document_title") or "")[:240],
                        section_path=str(hit.get("section_path") or "")[:512],
                        clause_id=str(hit.get("clause_id") or "")[:128],
                        support_type="direct",
                        confidence="high",
                        citation_json=json.dumps(
                            {"citation": {"field": "risk_escalation", "sentence": risk_sentence}, "source": hit},
                            ensure_ascii=False,
                            default=str,
                        ),
                    ))
            db.flush()
        except Exception as exc:
            logger.warning("报告知识引用审计保存失败: report_id={}, error={}", getattr(report, "id", None), exc)

    @staticmethod
    def integer_or_none(value: Any) -> Optional[int]:
        try:
            if value in (None, ""):
                return None
            return int(value)
        except (TypeError, ValueError):
            return None

    def select_llm_report(self, workflow_payload: dict[str, Any]) -> Optional[dict[str, Any]]:
        execution = workflow_payload.get("execution_result")
        if not isinstance(execution, dict):
            return None
        node_results = execution.get("node_results") or []
        if not isinstance(node_results, list):
            return None

        cloud_error = self.find_cloud_error(node_results) or workflow_payload.get("execution_error")
        cloud_report = self.find_node_report(node_results, preferred_ids={"action_report"})
        if cloud_report:
            cloud_report["cloud_error"] = cloud_error
            return cloud_report

        local_report = self.find_node_report(node_results, preferred_ids={"action_reasoning"})
        if local_report:
            local_report["source"] = "qwen4b"
            local_report["source_label"] = "Qwen3-VL-4B 本地场景理解"
            local_report["cloud_error"] = cloud_error
            return local_report
        return None

    def find_node_report(
        self,
        node_results: list[Any],
        *,
        preferred_ids: set[str],
    ) -> Optional[dict[str, Any]]:
        for row in node_results:
            if not isinstance(row, dict):
                continue
            node_id = str(row.get("node_id") or row.get("id") or "")
            if node_id not in preferred_ids:
                continue
            if str(row.get("status") or "").lower() != "success":
                continue
            output = row.get("output")
            text = self.summarize_structured_output(output) or self.extract_text(output)
            text = self.clean_model_output_field(text, allow_empty=True)
            if not text or self.looks_like_model_thinking(text):
                continue
            if node_id == "action_report":
                if not self.is_complete_cloud_report(output):
                    continue
                return {
                    "source": "qwen35b",
                    "source_label": "Qwen3.6-35B-A3B 云端增强分析",
                    "text": text,
                    "raw_output": output,
                    "node_id": node_id,
                }
            return {
                "source": "qwen4b",
                "source_label": "Qwen3-VL-4B 本地场景理解",
                "text": text,
                "raw_output": output,
                "node_id": node_id,
            }
        return None

    def is_complete_cloud_report(self, output: Any) -> bool:
        """Do not select a successful transport JSON as a report.

        Older workflow records may contain a supplemental runtime-state JSON
        under the cloud node even though the node was marked successful.
        Require the report fields before allowing it to replace the 4B result.
        """
        result = output.get("inference_result") if isinstance(output, dict) else output
        result = result if isinstance(result, dict) else {}
        candidate = result.get("final_report") if isinstance(result.get("final_report"), dict) else result
        required = (
            "detailed_scene_analysis",
            "risk_reasoning",
            "impact_assessment",
            "response_plan",
            "monitoring_suggestions",
        )
        present = sum(
            1
            for key in required
            if isinstance(candidate.get(key), str)
            and candidate.get(key).strip() not in {"", "—", "-"}
        )
        risk_level = str(result.get("risk_level") or candidate.get("risk_level") or "").lower()
        return present >= 3 and risk_level in {"low", "medium", "high", "critical"}

    def find_cloud_error(self, node_results: list[Any]) -> Optional[str]:
        for row in node_results:
            if not isinstance(row, dict):
                continue
            if str(row.get("node_id") or "") != "action_report":
                continue
            status = str(row.get("status") or "").lower()
            if status == "success":
                return None
            error = row.get("error") or row.get("message")
            if error:
                return str(error)
            output = row.get("output")
            if isinstance(output, dict) and output.get("error"):
                return str(output.get("error"))
        return None

    def extract_text(self, value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, str):
            return value.strip()
        if isinstance(value, dict):
            for key in ("response", "report", "analysis", "summary", "content", "text", "final_report"):
                text = self.extract_text(value.get(key))
                if text:
                    return text
            inference = value.get("inference_result")
            text = self.extract_text(inference)
            if text:
                return text
            choices = value.get("choices")
            if isinstance(choices, list):
                for choice in choices:
                    text = self.extract_text(choice)
                    if text:
                        return text
            message = value.get("message")
            text = self.extract_text(message)
            if text:
                return text
        if isinstance(value, list):
            for item in value:
                text = self.extract_text(item)
                if text:
                    return text
        return ""

    def summarize_structured_output(self, value: Any) -> str:
        fields = self.extract_template_overrides(value)
        if fields.get("handling_summary"):
            return str(fields["handling_summary"]).strip()
        if fields.get("summary"):
            return str(fields["summary"]).strip()
        if fields.get("conclusion"):
            return str(fields["conclusion"]).strip()
        if not isinstance(value, dict):
            return ""
        parts = []
        for key in (
            "scene_description",
            "suspected_event",
            "risk_level",
            "confidence",
            "evidence",
            "uncertainties",
            "recommendations",
            "risk_assessment",
            "emergency_suggestion",
        ):
            current = value.get(key)
            if current in (None, "", []):
                continue
            label = {
                "scene_description": "现场描述",
                "suspected_event": "疑似事件",
                "risk_level": "风险等级",
                "confidence": "置信度",
                "evidence": "证据",
                "uncertainties": "不确定因素",
                "recommendations": "处置建议",
                "risk_assessment": "风险研判",
                "emergency_suggestion": "应急建议",
            }.get(key, key)
            if isinstance(current, list):
                current = "；".join(str(item) for item in current if item)
            parts.append(f"{label}：{current}")
        return "\n".join(parts)

    def extract_template_overrides(self, value: Any) -> dict[str, Any]:
        candidates: list[dict[str, Any]] = []
        self.collect_template_dicts(value, candidates)
        overrides: dict[str, Any] = {}
        for candidate in candidates:
            explicit_summary = self.valid_report_text(candidate.get("handling_summary"))
            if explicit_summary:
                overrides["handling_summary"] = explicit_summary
            else:
                detailed = self.detailed_fields_summary(candidate)
                if detailed and not overrides.get("handling_summary"):
                    overrides["handling_summary"] = detailed
            for field in TEMPLATE_FIELDS:
                current = candidate.get(field)
                if current in (None, "", []):
                    continue
                if field == "handling_summary" and overrides.get("handling_summary"):
                    continue
                if isinstance(current, (dict, list)):
                    current = self.format_structured_value(current)
                current = self.clean_model_output_field(current, allow_empty=True)
                if current in (None, "", "—"):
                    continue
                overrides[field] = current
            if not overrides.get("key_observation"):
                risk_reasoning = candidate.get("risk_reasoning")
                if isinstance(risk_reasoning, str) and risk_reasoning.strip():
                    cleaned = self.clean_model_output_field(risk_reasoning, allow_empty=True)
                    if cleaned not in ("", "—"):
                        overrides["key_observation"] = self.compact(cleaned, 260)
            if not overrides.get("conclusion"):
                conclusion = candidate.get("impact_assessment") or candidate.get("monitoring_suggestions")
                if isinstance(conclusion, str) and conclusion.strip():
                    cleaned = self.clean_model_output_field(conclusion, allow_empty=True)
                    if cleaned not in ("", "—"):
                        overrides["conclusion"] = self.compact(cleaned, 320)
        return overrides

    def detailed_fields_summary(self, value: dict[str, Any]) -> str:
        fields = [
            ("一、现场场景", value.get("detailed_scene_analysis")),
            ("二、风险研判", value.get("risk_reasoning")),
            ("三、影响评估", value.get("impact_assessment")),
            ("四、处置建议", value.get("response_plan")),
            ("五、持续监测", value.get("monitoring_suggestions")),
        ]
        lines = [f"{label}：{text}" for label, text in fields if isinstance(text, str) and text.strip()]
        return "\n".join(lines)

    def collect_template_dicts(self, value: Any, candidates: list[dict[str, Any]]) -> None:
        if isinstance(value, dict):
            if any(key in value for key in TEMPLATE_FIELDS):
                candidates.append(value)
            for key in ("template_fields", "template_data", "docx_context", "final_report", "inference_result"):
                current = value.get(key)
                if isinstance(current, dict):
                    self.collect_template_dicts(current, candidates)
            choices = value.get("choices")
            if isinstance(choices, list):
                for choice in choices:
                    self.collect_template_dicts(choice, candidates)
        elif isinstance(value, list):
            for item in value:
                self.collect_template_dicts(item, candidates)

    def format_structured_value(self, value: Any) -> str:
        if isinstance(value, list):
            return "；".join(str(item) for item in value if item)
        if isinstance(value, dict):
            return "；".join(f"{key}：{val}" for key, val in value.items() if val not in (None, "", []))
        return str(value)

    def build_context(
        self,
        db: Session,
        instance: SafetyEventInstance,
        event: EventLibrary,
        workflow_payload: dict[str, Any],
        selected: dict[str, Any],
    ) -> dict[str, Any]:
        now = dt.datetime.now(dt.timezone.utc)
        report_date = self.to_local_datetime(instance.started_at or now).date()
        source = db.query(DataSource).filter(DataSource.id == instance.data_source_id).first()
        timeline = (
            db.query(SafetyEventTimelineLog)
            .options(
                load_only(
                    SafetyEventTimelineLog.id,
                    SafetyEventTimelineLog.event_instance_id,
                    SafetyEventTimelineLog.log_type,
                    SafetyEventTimelineLog.status,
                    SafetyEventTimelineLog.message,
                    SafetyEventTimelineLog.payload,
                    SafetyEventTimelineLog.create_time,
                )
            )
            .filter(SafetyEventTimelineLog.event_instance_id == instance.id)
            .order_by(SafetyEventTimelineLog.id.asc())
            .all()
        )
        evidence = (
            db.query(SafetyEventEvidence)
            .filter(SafetyEventEvidence.event_instance_id == instance.id)
            .order_by(SafetyEventEvidence.captured_at.asc(), SafetyEventEvidence.id.asc())
            .all()
        )
        visual = self.visual_snapshot(instance, timeline)
        camera = self.find_camera(db, instance, visual)
        # 联动设备一次任务可能归档多张取证图（无人机、机器狗、人工处置）。
        # 报告对每一种实际执行且有图像回传的联动对象各展示一张代表图，
        # 同一对象的其余图片仍仅归档，避免一次巡航的多帧挤占报告版面。
        model_image_items = self.collect_image_items(workflow_payload, visual, evidence)
        video_items = self.collect_video_items(workflow_payload, visual, evidence)
        if not model_image_items and video_items:
            model_image_items = self.extract_frame_items_from_videos(
                video_items,
                event_name=getattr(event, "event_name", None) or instance.summary or "安全事件",
                analysis_text=self.clean_report_text(selected["text"]),
                workflow_insight={},
            )
        model_report_image = self.select_model_report_image(model_image_items)
        # 云端工作流媒体地址可能过期。先由 ``select_model_report_image`` 映射到
        # 同一记录里的本地 MinIO 对象；若本地对象也无法读取，则从已归档的事件
        # 视频取一帧，保证报告不会只剩联动设备的图片。
        if (
            (not model_report_image or not self.is_readable_image_item(model_report_image))
            and video_items
        ):
            fallback_items = self.extract_frame_items_from_videos(
                video_items,
                event_name=getattr(event, "event_name", None) or instance.summary or "安全事件",
                analysis_text=self.clean_report_text(selected["text"]),
                workflow_insight={},
            )
            fallback_image = self.select_model_report_image(fallback_items)
            if fallback_image:
                model_report_image = fallback_image
        linkage_report_images = self.select_linkage_report_images(evidence)
        image_items = [item for item in (model_report_image, *linkage_report_images) if item]
        selected_text = self.clean_report_text(selected["text"])
        workflow_insight = self.workflow_insight(workflow_payload, visual, selected_text)
        report_sources, report_citations = self.report_knowledge_citations(
            workflow_payload,
            selected,
            workflow_insight,
        )
        # Keep the all-node values available for audit/debugging, while the
        # rendered report and its citation table use only the final node's
        # citations that are actually present in its report text.
        workflow_insight["knowledge_sources_all"] = workflow_insight.get("knowledge_sources") or []
        workflow_insight["sentence_citations_all"] = workflow_insight.get("sentence_citations") or []
        workflow_insight["knowledge_sources"] = report_sources
        workflow_insight["sentence_citations"] = report_citations
        workflow_insight["knowledge_sources_summary"] = self.format_knowledge_sources(
            report_sources,
            report_citations,
        )
        # The final report may combine the edge model's citations with the
        # cloud review. Keep one shared source order for every rendered field;
        # otherwise the cloud response numbers its own subset independently
        # from the report's merged knowledge-basis section.
        merged_citation_sources = workflow_insight.get("knowledge_sources")
        if isinstance(merged_citation_sources, list):
            selected["citation_sources"] = merged_citation_sources
        merged_sentence_citations = workflow_insight.get("sentence_citations")
        if isinstance(merged_sentence_citations, list):
            selected["citation_ids"] = [
                str(evidence_id).strip()
                for citation in merged_sentence_citations
                if isinstance(citation, dict)
                for evidence_id in citation.get("evidence_ids") or []
                if str(evidence_id).strip()
            ]
        observation = dict(instance.latest_observation or {})
        risk_escalation = observation.get("risk_escalation") if isinstance(observation.get("risk_escalation"), dict) else {}
        supplemental_context = observation.get("supplemental_context") if isinstance(observation.get("supplemental_context"), dict) else {}
        if risk_escalation:
            workflow_insight["risk_escalation"] = risk_escalation
            workflow_insight["risk_escalation_summary"] = self.risk_escalation_summary(risk_escalation, supplemental_context)
            # Keep escalation hits in the audit context, but do not render
            # them as final report citations unless the selected model cites
            # them in the report body.
            workflow_insight["risk_escalation_knowledge_summary"] = self.risk_escalation_knowledge_summary(risk_escalation)
        workflow_insight["event_name"] = getattr(event, "event_name", None) or instance.summary or "安全事件"
        workflow_insight["event_code"] = getattr(event, "event_code", None)
        timeline_summary = self.timeline_summary(timeline)
        evidence_image_path = self.select_evidence_image(image_items, video_items)

        source_label = self.source_label(instance, source, selected)
        location_text = self.location(source, camera, visual)
        cloud_note = ""
        if selected.get("source") == "qwen4b" and selected.get("cloud_error"):
            cloud_note = "云端增强暂不可用，本报告已采用本地 4B 分析结果生成。"
        risk_escalation_summary = str(workflow_insight.get("risk_escalation_summary") or "").strip()
        risk_assessment_detail = (
            risk_escalation_summary
            or self.final_report_field(selected, "risk_reasoning", workflow_insight.get("raw_excerpt"))
        )

        context = {
            "workflow_insight": workflow_insight,
            "report_date": report_date,
            "report_date_cn": self.format_chinese_date(report_date),
            "report_time": self.format_datetime(dt.datetime.now(LOCAL_TIMEZONE)),
            "event_name": getattr(event, "event_name", None) or instance.summary or "安全事件",
            "instance_no": instance.instance_no,
            "instance_no_prefix": self.instance_no_parts(instance.instance_no)[0],
            "instance_no_suffix": self.instance_no_parts(instance.instance_no)[1],
            "risk_label": self.report_risk_label(instance, selected),
            "result_label": self.result_label(instance),
            "occur_time": self.format_datetime(instance.started_at),
            "occur_time_display": self.format_datetime(instance.started_at).replace(" ", "\n"),
            "completed_at": self.completed_at(instance, timeline),
            "event_duration": self.event_duration(instance, timeline),
            "emergency_level": self.emergency_level(instance),
            "confidence_label": self.confidence_label(selected),
            "source_label": source_label,
            "source_trigger_label": self.source_trigger_label(instance, source),
            "location": location_text,
            "location_short": self.compact(location_text, 36),
            "evidence_count": len(image_items),
            "summary": self.event_summary(instance, event, source, visual, workflow_insight),
            "key_observation": self.key_observation(
                workflow_payload,
                visual,
                selected_text,
                workflow_insight,
                selected,
            ),
            "source_summary": self.source_summary(image_items, video_items, selected),
            "handling_source": selected["source_label"],
            "timeline_count": len(timeline),
            "trigger_summary": self.trigger_summary(instance, event, visual, workflow_insight),
            "screening_summary": self.screening_summary(visual),
            "model_route_summary": self.model_route_summary(
                workflow_payload,
                selected,
                getattr(event, "event_name", None) or instance.summary or "",
            ),
            "workflow_nodes_summary": self.workflow_nodes_summary(workflow_payload),
            "specialized_summary": self.specialized_summary(workflow_insight),
            "local_analysis_summary": self.node_analysis_summary(workflow_payload, "action_reasoning"),
            "cloud_analysis_summary": self.node_analysis_summary(workflow_payload, "action_report"),
            "scene_detail": self.final_report_field(selected, "detailed_scene_analysis", selected_text),
            "risk_assessment_detail": risk_assessment_detail,
            "impact_assessment": self.final_report_field(selected, "impact_assessment", "需结合现场巡查、水位雨量和坝体状态持续确认影响范围。"),
            "response_plan": self.final_report_field(selected, "response_plan", "维持事件取证和现场复核，按风险等级启动相应联动处置。"),
            "monitoring_suggestions": self.final_report_field(selected, "monitoring_suggestions", "持续跟踪摄像头画面、水位、雨量、风速和坝体安全监测数据。"),
            "recommendations_text": self.recommendations_text(selected),
            "evidence_inventory": self.evidence_inventory(image_items, video_items),
            "frame_evidence_summary": self.frame_evidence_summary(image_items, video_items),
            "linkage_evidence_summary": self.linkage_evidence_summary(evidence),
            "analysis_limitations": self.analysis_limitations(selected, workflow_insight, image_items, video_items),
            "follow_up_actions": self.follow_up_actions(instance, selected),
            "handling_summary": self.handling_summary(
                instance=instance,
                event=event,
                visual=visual,
                selected=selected,
                selected_text=selected_text,
                workflow_insight=workflow_insight,
                image_items=image_items,
                video_items=video_items,
            ),
            "timeline_summary": timeline_summary,
            "evidence_summary": self.evidence_summary(image_items, video_items),
            "evidence_caption": self.evidence_caption(image_items, video_items),
            "evidence_image": evidence_image_path,
            "conclusion": self.build_conclusion(selected, workflow_insight, cloud_note),
        }
        overrides = self.extract_template_overrides(selected.get("raw_output"))
        context.update({
            key: value
            for key, value in overrides.items()
            if key not in SYSTEM_FACT_FIELDS
        })
        context["report_date"] = report_date
        context["report_date_cn"] = self.format_chinese_date(report_date)
        context["report_time"] = self.format_datetime(dt.datetime.now(LOCAL_TIMEZONE))
        context["event_name"] = getattr(event, "event_name", None) or instance.summary or "安全事件"
        context["instance_no"] = instance.instance_no
        context["instance_no_prefix"], context["instance_no_suffix"] = self.instance_no_parts(instance.instance_no)
        context["risk_label"] = self.report_risk_label(instance, selected)
        context["result_label"] = self.result_label(instance)
        context["occur_time"] = self.format_datetime(instance.started_at)
        context["occur_time_display"] = self.format_datetime(instance.started_at).replace(" ", "\n")
        context["completed_at"] = self.completed_at(instance, timeline)
        context["event_duration"] = self.event_duration(instance, timeline)
        context["evidence_count"] = len(image_items)
        context["timeline_count"] = len(timeline)
        context["timeline_summary"] = timeline_summary
        context["evidence_caption"] = self.evidence_caption(image_items, video_items)
        context["evidence_image"] = evidence_image_path
        if risk_escalation_summary:
            context["risk_assessment_detail"] = risk_assessment_detail
            context.update(self.risk_escalation_report_overrides(risk_escalation_summary, supplemental_context))
        if selected.get("source") == "qwen4b":
            if not context.get("handling_summary"):
                context["handling_summary"] = self.handling_summary(
                    instance=instance,
                    event=event,
                    visual=visual,
                    selected=selected,
                    selected_text=selected_text,
                    workflow_insight=workflow_insight,
                    image_items=image_items,
                    video_items=video_items,
                )
            if not risk_escalation_summary:
                context["conclusion"] = self.build_conclusion(selected, workflow_insight, cloud_note)
        else:
            if not risk_escalation_summary:
                context["conclusion"] = self.build_conclusion(selected, workflow_insight, cloud_note)
        self.normalize_report_context_citations(context, selected)
        return context

    def render_docx(self, context: dict[str, Any]) -> bytes:
        template = DocxTemplate(str(TEMPLATE_PATH))
        image_path = context.pop("evidence_image", None)
        render_context = dict(context)
        render_context.setdefault("report_date_cn", self.format_chinese_date(render_context.get("report_date")))
        if image_path:
            render_context["evidence_image"] = InlineImage(template, str(image_path), width=Mm(130))
        else:
            render_context["evidence_image"] = "未获取到可嵌入图像"
        template.render(render_context)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            temp_path = Path(tmp.name)
        try:
            template.save(str(temp_path))
            self.append_knowledge_section(
                temp_path,
                (context.get("workflow_insight") or {}).get("knowledge_sources_summary")
                or context.get("knowledge_sources_summary"),
            )
            self.normalize_docx_knowledge_refs(
                temp_path,
                (context.get("workflow_insight") or {}).get("knowledge_sources_summary")
                or context.get("knowledge_sources_summary"),
            )
            self.rebuild_toc_bookmark_links(temp_path)
            self.disable_docx_proofing(temp_path)
            self.normalize_docx_fonts(temp_path)
            self.normalize_first_page_date_footer(temp_path)
            self.refresh_toc_page_number_fallbacks(temp_path)
            return temp_path.read_bytes()
        finally:
            try:
                temp_path.unlink(missing_ok=True)
            except OSError:
                pass
            if image_path:
                try:
                    Path(image_path).unlink(missing_ok=True)
                except OSError:
                    pass

    def append_knowledge_section(self, docx_path: Path, knowledge_summary: Any) -> None:
        text = str(knowledge_summary or "").strip()
        if not text:
            return
        try:
            document = Document(str(docx_path))
            heading_exists = any(
                paragraph.style.name.startswith("Heading")
                and re.sub(r"\s+", "", paragraph.text) in {"6知识依据", "06知识依据"}
                for paragraph in document.paragraphs
            )
            if heading_exists:
                return
            heading = document.add_paragraph()
            heading.style = "Heading 1"
            heading.paragraph_format.space_before = Pt(14)
            heading.paragraph_format.space_after = Pt(8)
            heading.paragraph_format.line_spacing = Pt(20)
            run = heading.add_run("6  知识依据")
            run.bold = True
            run.font.size = Pt(15)
            run.font.name = "经典宋体简"
            run.font.color.rgb = RGBColor(0, 0, 0)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "经典宋体简")
            run._element.rPr.append(OxmlElement("w:noProof"))

            for run in heading.runs:
                run.bold = True
            for line in text.splitlines():
                cleaned = self.normalize_punctuation(line.strip())
                if not cleaned:
                    continue
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)
                run = paragraph.add_run(cleaned)
                run.font.size = Pt(11)
                run.font.name = "经典宋体简"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "经典宋体简")
                run._element.rPr.append(OxmlElement("w:noProof"))
            document.save(str(docx_path))
        except Exception as exc:
            logger.warning("追加报告知识依据章节失败: {}", exc)

    def normalize_docx_knowledge_refs(self, docx_path: Path, knowledge_summary: Any) -> None:
        summary = str(knowledge_summary or "").strip()
        match = re.search(r"\[(\d+)\]", summary)
        first_index = match.group(1) if match else ("1" if summary else "")
        clause_index_map = self.knowledge_summary_clause_index_map(summary)
        try:
            document = Document(str(docx_path))
            changed = False

            def normalize_runs(paragraph) -> None:
                nonlocal changed
                for run in paragraph.runs:
                    original = run.text
                    if not original:
                        continue
                    updated = original
                    if first_index:
                        updated = re.sub(r"\[K\d+\]", f"[{first_index}]", updated)
                        updated = re.sub(r"(?<!\[)K\d+(?!\])", f"[{first_index}]", updated)
                        updated = re.sub(r"知识库依据(?!\s*\[\d+\])", f"知识库依据[{first_index}]", updated)
                        updated = re.sub(r"结合知识库(?!依据)", f"结合知识库依据[{first_index}]", updated)
                    for clause_id, index in clause_index_map.items():
                        updated = updated.replace(f"[{clause_id}]", f"[{index}]")
                        updated = re.sub(
                            rf"(?<![A-Za-z0-9_\[]){re.escape(clause_id)}(?![A-Za-z0-9_\]])",
                            f"[{index}]",
                            updated,
                        )
                    updated = self.normalize_punctuation(updated)
                    if updated != original:
                        run.text = updated
                        changed = True

            in_knowledge_section = False
            for paragraph in document.paragraphs:
                normalized_heading = re.sub(r"\s+", "", paragraph.text or "")
                style_name = paragraph.style.name if paragraph.style is not None else ""
                if style_name.startswith("Heading") and normalized_heading in {"6知识依据", "06知识依据"}:
                    in_knowledge_section = True
                    continue
                if in_knowledge_section:
                    continue
                normalize_runs(paragraph)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            normalize_runs(paragraph)
            if changed:
                document.save(str(docx_path))
        except Exception as exc:
            logger.warning("清洗报告知识引用编号失败: {}", exc)

    @staticmethod
    def knowledge_summary_clause_index_map(summary: str) -> dict[str, str]:
        mapping: dict[str, str] = {}
        for line in str(summary or "").splitlines():
            index_match = re.search(r"\[(\d+)\]", line)
            if not index_match:
                continue
            index = index_match.group(1)
            for clause in re.findall(r"\b[A-Z][A-Z0-9_-]*-\d+\b", line):
                mapping.setdefault(clause, index)
            for evidence in re.findall(r"\bK\d+\b", line):
                mapping.setdefault(evidence, index)
        return mapping

    def rebuild_toc_bookmark_links(self, docx_path: Path) -> None:
        entries = [
            ("1  事件信息", "dam_report_heading_1"),
            ("2  事件复核结果", "dam_report_heading_2"),
            ("3  研判与处置记录", "dam_report_heading_3"),
            ("4  证据材料", "dam_report_heading_4"),
            ("5  处置结论", "dam_report_heading_5"),
            ("6  知识依据", "dam_report_heading_6"),
        ]
        normalized_entries = {
            re.sub(r"\s+", "", title): (title, bookmark)
            for title, bookmark in entries
        }
        try:
            document = Document(str(docx_path))
            bookmark_id = 6100
            for paragraph in document.paragraphs:
                normalized = re.sub(r"\s+", "", paragraph.text or "")
                if normalized not in normalized_entries:
                    continue
                _, bookmark_name = normalized_entries[normalized]
                self.ensure_paragraph_bookmark(
                    self.bookmark_offset_target(document, paragraph) or paragraph,
                    bookmark_name,
                    bookmark_id,
                )
                bookmark_id += 1

            toc_started = False
            for paragraph in document.paragraphs:
                if paragraph.text.strip() == "目录":
                    toc_started = True
                    continue
                if not toc_started:
                    continue
                normalized = re.sub(r"\s+", "", (paragraph.text or "").split("\t", 1)[0])
                normalized = re.sub(r"\d+$", "", normalized)
                if normalized in normalized_entries:
                    title, bookmark_name = normalized_entries[normalized]
                    self.rewrite_toc_paragraph(paragraph, title, bookmark_name)
                    continue
                if paragraph.text.strip() == "":
                    break

            settings = document.settings._element
            update_fields = settings.find(qn("w:updateFields"))
            if update_fields is None:
                update_fields = OxmlElement("w:updateFields")
                settings.append(update_fields)
            update_fields.set(qn("w:val"), "true")
            document.save(str(docx_path))
        except Exception as exc:
            logger.warning("重建报告目录跳转链接失败: {}", exc)

    def refresh_toc_page_number_fallbacks(self, docx_path: Path) -> None:
        """Write actual heading page numbers into TOC field fallback text.

        OnlyOffice may not refresh PAGEREF field results immediately when opening a
        generated DOCX. The fields are still kept for jump/update support, while
        the visible fallback text is updated from a temporary PDF rendering.
        """
        try:
            page_numbers = self.detect_heading_page_numbers(docx_path)
            if not page_numbers:
                return
            document_xml = "word/document.xml"
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                temp_zip_path = Path(tmp.name)
            try:
                changed = False
                with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
                    temp_zip_path,
                    "w",
                    zipfile.ZIP_DEFLATED,
                ) as zout:
                    for item in zin.infolist():
                        data = zin.read(item.filename)
                        if item.filename == document_xml:
                            xml = data.decode("utf-8")
                            for bookmark_name, page_number in page_numbers.items():
                                pattern = (
                                    r"(<w:instrText\b[^>]*>\s*PAGEREF\s+"
                                    + re.escape(bookmark_name)
                                    + r"\s+\\h\s*</w:instrText>\s*"
                                    + r"<w:fldChar\b[^>]*w:fldCharType=\"separate\"[^>]*/>\s*"
                                    + r"<w:t>)(.*?)(</w:t>)"
                                )
                                xml, count = re.subn(
                                    pattern,
                                    lambda match, number=str(page_number): f"{match.group(1)}{number}{match.group(3)}",
                                    xml,
                                    count=1,
                                    flags=re.DOTALL,
                                )
                                changed = changed or count > 0
                            data = xml.encode("utf-8")
                        zout.writestr(item, data)
                if changed:
                    temp_zip_path.replace(docx_path)
                else:
                    temp_zip_path.unlink(missing_ok=True)
            finally:
                temp_zip_path.unlink(missing_ok=True)
        except Exception as exc:
            logger.warning("刷新报告目录页码失败: {}", exc)

    def detect_heading_page_numbers(self, docx_path: Path) -> dict[str, int]:
        try:
            from pypdf import PdfReader
        except Exception as exc:
            logger.warning("刷新目录页码跳过：缺少 PDF 文本解析库 pypdf: {}", exc)
            return {}

        headings = {
            "dam_report_heading_1": "1事件信息",
            "dam_report_heading_2": "2事件复核结果",
            "dam_report_heading_3": "3研判与处置记录",
            "dam_report_heading_4": "4证据材料",
            "dam_report_heading_5": "5处置结论",
            "dam_report_heading_6": "6知识依据",
        }
        with tempfile.TemporaryDirectory(prefix="dam_report_toc_pages_") as temp_dir:
            work_dir = Path(temp_dir)
            output_dir = work_dir / "output"
            profile_dir = work_dir / "profile"
            runtime_dir = work_dir / "runtime"
            for directory in (output_dir, profile_dir, runtime_dir):
                directory.mkdir(parents=True, exist_ok=True)
            runtime_dir.chmod(0o700)
            environment = os.environ.copy()
            environment["XDG_RUNTIME_DIR"] = str(runtime_dir)
            command = [
                "libreoffice",
                "--headless",
                f"-env:UserInstallation={profile_dir.as_uri()}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(docx_path),
            ]
            result = subprocess.run(
                command,
                cwd=work_dir,
                env=environment,
                capture_output=True,
                text=True,
                timeout=90,
                check=False,
            )
            output_files = list(output_dir.glob("*.pdf"))
            if result.returncode != 0 or not output_files:
                detail = (result.stderr or result.stdout or "").strip()
                logger.warning("刷新目录页码跳过：DOCX 转 PDF 失败 {}", detail[-240:])
                return {}

            reader = PdfReader(str(output_files[0]))
            found: dict[str, int] = {}
            for page_index, page in enumerate(reader.pages, start=1):
                # Cover and TOC contain the same titles; real section headings
                # start after them in the current report template.
                if page_index <= 2:
                    continue
                try:
                    text = page.extract_text() or ""
                except Exception:
                    continue
                normalized = re.sub(r"\s+", "", text)
                for bookmark_name, heading_text in headings.items():
                    if bookmark_name not in found and heading_text in normalized:
                        found[bookmark_name] = page_index
                if len(found) == len(headings):
                    break
            return found

    @staticmethod
    def bookmark_offset_target(document: Document, paragraph):
        fallback = None
        non_heading_paragraphs = 0
        sibling = paragraph._p.getnext()
        while sibling is not None:
            if sibling.tag == qn("w:p"):
                text = "".join(node.text or "" for node in sibling.iter(qn("w:t"))).strip()
                if text:
                    for candidate in document.paragraphs:
                        if candidate._p is sibling:
                            if candidate.style.name.startswith("Heading"):
                                return fallback
                            fallback = fallback or candidate
                            non_heading_paragraphs += 1
                            if non_heading_paragraphs >= 2:
                                return candidate
            elif sibling.tag == qn("w:tbl"):
                for table in document.tables:
                    if table._tbl is sibling and table.rows and table.rows[0].cells:
                        first_cell = table.rows[0].cells[0]
                        return first_cell.paragraphs[0] if first_cell.paragraphs else None
            sibling = sibling.getnext()
        return fallback

    @staticmethod
    def ensure_paragraph_bookmark(paragraph, bookmark_name: str, bookmark_id: int) -> None:
        root = paragraph._p.getroottree().getroot()
        for existing in list(root.iter(qn("w:bookmarkStart"))):
            if existing.get(qn("w:name")) == bookmark_name:
                parent = existing.getparent()
                if parent is not None:
                    parent.remove(existing)
        for existing in list(root.iter(qn("w:bookmarkEnd"))):
            if existing.get(qn("w:id")) == str(bookmark_id):
                parent = existing.getparent()
                if parent is not None:
                    parent.remove(existing)
        paragraph_element = paragraph._p
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), bookmark_name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))
        paragraph_element.append(start)
        paragraph_element.append(end)

    def rewrite_toc_paragraph(self, paragraph, title: str, bookmark_name: str) -> None:
        paragraph_element = paragraph._p
        p_pr = paragraph_element.get_or_add_pPr()
        tabs = p_pr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            p_pr.append(tabs)
        for child in list(tabs):
            tabs.remove(child)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "8787")
        tabs.append(tab)
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "520")
        spacing.set(qn("w:lineRule"), "exact")

        for child in list(paragraph_element):
            if child.tag != qn("w:pPr"):
                paragraph_element.remove(child)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), bookmark_name)
        hyperlink.set(qn("w:history"), "1")
        run = OxmlElement("w:r")
        run.append(self.run_properties_xml(font="经典宋体简", size_half_points=30))
        text = OxmlElement("w:t")
        text.text = title
        run.append(text)
        hyperlink.append(run)
        paragraph_element.append(hyperlink)

        tab_run = OxmlElement("w:r")
        tab_run.append(self.run_properties_xml(font="经典宋体简", size_half_points=30))
        tab_run.append(OxmlElement("w:tab"))
        paragraph_element.append(tab_run)

        field_run = OxmlElement("w:r")
        field_run.append(self.run_properties_xml(font="经典宋体简", size_half_points=30))
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        begin.set(qn("w:dirty"), "true")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" PAGEREF {bookmark_name} \\h "
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        fallback = OxmlElement("w:t")
        fallback.text = "1"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        field_run.extend([begin, instr, separate, fallback, end])
        paragraph_element.append(field_run)

    @staticmethod
    def run_properties_xml(font: str, size_half_points: int) -> OxmlElement:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Times New Roman")
        r_fonts.set(qn("w:hAnsi"), "Times New Roman")
        r_fonts.set(qn("w:eastAsia"), font)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), str(size_half_points))
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "none")
        no_proof = OxmlElement("w:noProof")
        r_pr.extend([r_fonts, color, size, underline, no_proof])
        return r_pr

    def disable_docx_proofing(self, docx_path: Path) -> None:
        try:
            document = Document(str(docx_path))
            settings = document.settings._element
            for tag in ("w:hideSpellingErrors", "w:hideGrammaticalErrors"):
                element = settings.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    settings.append(element)
                element.set(qn("w:val"), "true")
            proof_state = settings.find(qn("w:proofState"))
            if proof_state is None:
                proof_state = OxmlElement("w:proofState")
                settings.append(proof_state)
            proof_state.set(qn("w:spelling"), "clean")
            proof_state.set(qn("w:grammar"), "clean")

            for paragraph in document.paragraphs:
                self.disable_paragraph_proofing(paragraph)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.disable_paragraph_proofing(paragraph)
            document.save(str(docx_path))
        except Exception as exc:
            logger.warning("关闭报告拼写检查标记失败: {}", exc)

    @staticmethod
    def disable_paragraph_proofing(paragraph) -> None:
        for run in paragraph.runs:
            r_pr = run._element.get_or_add_rPr()
            if r_pr.find(qn("w:noProof")) is None:
                r_pr.append(OxmlElement("w:noProof"))

    def normalize_docx_fonts(self, docx_path: Path) -> None:
        try:
            document = Document(str(docx_path))
            self.set_style_font(document.styles["Normal"], "经典宋体简", size=Pt(11))
            self.set_style_font(document.styles["Heading 1"], "黑体", size=Pt(15), bold=True)

            for index, paragraph in enumerate(document.paragraphs):
                self.normalize_paragraph_font(paragraph, paragraph_index=index)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.normalize_paragraph_font(paragraph, in_table=True)
            for section in document.sections:
                for part in (
                    section.header,
                    section.first_page_header,
                    section.footer,
                    section.first_page_footer,
                ):
                    for paragraph in part.paragraphs:
                        self.normalize_paragraph_font(paragraph)
                    for table in part.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    self.normalize_paragraph_font(paragraph, in_table=True)
            document.save(str(docx_path))
        except Exception as exc:
            logger.warning("统一报告字体失败: {}", exc)

    def normalize_first_page_date_footer(self, docx_path: Path) -> None:
        try:
            document = Document(str(docx_path))
            for section in document.sections:
                for paragraph in section.first_page_footer.paragraphs:
                    if not re.search(r"\d{4}年\d{1,2}月\d{1,2}日", paragraph.text or ""):
                        continue
                    for run in paragraph.runs:
                        self.set_run_font(run, "黑体", size_pt=22, bold=True)
                        run.font.name = "黑体"
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        r_pr = run._element.get_or_add_rPr()
                        r_fonts = r_pr.find(qn("w:rFonts"))
                        if r_fonts is None:
                            r_fonts = OxmlElement("w:rFonts")
                            r_pr.append(r_fonts)
                        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                            r_fonts.set(qn(key), "黑体")
            document.save(str(docx_path))
        except Exception as exc:
            logger.warning("统一报告首页日期字体失败: {}", exc)

    @staticmethod
    def set_style_font(style, font_name: str, *, size: Optional[Pt] = None, bold: Optional[bool] = None) -> None:
        style.font.name = font_name
        if size is not None:
            style.font.size = size
        if bold is not None:
            style.font.bold = bold
        r_pr = style._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(key), "Times New Roman" if key in {"w:ascii", "w:hAnsi"} else font_name)

    def normalize_paragraph_font(self, paragraph, *, in_table: bool = False, paragraph_index: Optional[int] = None) -> None:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if paragraph_index == 7:
            for run in paragraph.runs:
                self.set_run_font(run, "黑体", size_pt=36, bold=True)
            return
        if paragraph_index == 8:
            for run in paragraph.runs:
                self.set_run_font(run, "经典宋体简", size_pt=20, bold=False)
            return
        if style_name.startswith("Heading"):
            for run in paragraph.runs:
                self.set_run_font(run, "黑体", size_pt=15, bold=True)
            return
        for run in paragraph.runs:
            if not run.text:
                continue
            bold = bool(run.bold)
            text = run.text.strip()
            is_label = bold or text.endswith("：") or text in {
                "事件来源", "证据数量", "发生位置", "事件类型", "发生时间", "完成时间",
                "风险等级", "处置状态", "模型链路", "处置耗时", "初筛结论", "联动记录",
                "代表性证据图像",
            }
            font_name = "黑体" if is_label else "经典宋体简"
            size_pt = 10.5 if in_table else (15 if "\t" in run.text else 11)
            self.set_run_font(run, font_name, size_pt=size_pt, bold=is_label)

    @staticmethod
    def set_run_font(run, font_name: str, *, size_pt: float, bold: bool) -> None:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:ascii"), "Times New Roman")
        r_fonts.set(qn("w:hAnsi"), "Times New Roman")
        r_fonts.set(qn("w:eastAsia"), font_name)
        r_fonts.set(qn("w:cs"), font_name)
        if r_pr.find(qn("w:noProof")) is None:
            r_pr.append(OxmlElement("w:noProof"))

    def upsert_analysis_report(
        self,
        db: Session,
        *,
        instance: SafetyEventInstance,
        event: EventLibrary,
        file_url: str,
        report_date: dt.date,
    ) -> AnalysisReport:
        report = None
        if instance.analysis_report_id:
            report = db.query(AnalysisReport).filter(AnalysisReport.id == instance.analysis_report_id).first()
        if not report:
            report_no = f"EVR_{instance.instance_no}"
            report = db.query(AnalysisReport).filter(AnalysisReport.report_no == report_no).first()
        if not report:
            report = AnalysisReport(
                report_no=f"EVR_{instance.instance_no}",
                report_title=f"{getattr(event, 'event_name', None) or '安全事件'}处置报告",
                report_type="event",
                report_date=report_date,
                file_url=file_url,
            )
            db.add(report)
            db.flush()
        else:
            report.report_title = f"{getattr(event, 'event_name', None) or '安全事件'}处置报告"
            report.report_type = "event"
            report.report_date = report_date
            report.file_url = file_url
        instance.analysis_report_id = report.id
        db.flush()
        return report

    def find_camera(
        self,
        db: Session,
        instance: SafetyEventInstance,
        visual: Optional[dict[str, Any]] = None,
    ) -> Optional[Camera]:
        visual = visual or self.visual_snapshot(instance)
        camera_id = visual.get("camera_id") or instance.source_id
        if camera_id and str(camera_id).isdigit():
            return db.query(Camera).filter(Camera.id == int(camera_id)).first()
        return None

    def visual_snapshot(
        self,
        instance: SafetyEventInstance,
        timeline: Optional[list[SafetyEventTimelineLog]] = None,
    ) -> dict[str, Any]:
        # Event reports must describe the trigger evidence, not the later
        # recovery frame that may say "no target".
        for row in timeline or []:
            if str(row.log_type or "").upper() != "TRIGGER":
                continue
            payload = row.payload if isinstance(row.payload, dict) else {}
            observation = payload.get("observation") if isinstance(payload.get("observation"), dict) else {}
            visual = observation.get("visual") if isinstance(observation.get("visual"), dict) else {}
            if visual:
                return dict(visual)

        observation = dict(instance.latest_observation or {})
        visual = observation.get("visual")
        return dict(visual) if isinstance(visual, dict) else {}

    def collect_image_items(
        self,
        workflow_payload: dict[str, Any],
        visual: dict[str, Any],
        evidence: list[SafetyEventEvidence],
    ) -> list[dict[str, Any]]:
        candidates: list[dict[str, Any]] = []
        # 正式报告证据图只使用 4B 从均匀抽取复核帧中选出的代表帧。
        for key in (
            "representative_frame",
            "representative_frame_candidates",
            "key_frames",
            "representative_frames",
            "media_objects",
            "image_urls",
            "images",
        ):
            self.extend_image_media_items(candidates, self.find_nested_values(workflow_payload, key))
        # 模型库工作流没有返回代表帧时，使用 4B 初筛阶段均匀抽取的中间帧。
        # 这不是联动设备图，必须与后续的机器狗/无人机代表取证图同时保留。
        screening_frames = visual.get("image_urls") if isinstance(visual.get("image_urls"), list) else []
        screening_frames = [str(item) for item in screening_frames if str(item or "").strip()]
        if screening_frames:
            candidates.append({
                "url": screening_frames[len(screening_frames) // 2],
                "caption": "4B 初筛代表性抽帧",
                "source": "qwen4b_camera_screening",
                "role": "model_representative",
            })
        for row in evidence:
            if str(row.evidence_type or "").upper() in {"IMAGE", "CAMERA_SNAPSHOT"}:
                candidates.append({"url": row.file_url, "caption": row.description or "事件图像"})
        items = [
            item
            for item in self.unique_media_items(candidates)
            if not self.is_yolo_detection_image(item, item.get("url"))
        ]
        items.sort(key=lambda item: self.image_media_priority(item))
        return items[:8]

    def select_model_report_image(self, items: list[dict[str, Any]]) -> Optional[dict[str, Any]]:
        """Pick one model-side frame for the report, independently of action evidence."""
        if not items:
            return None
        qwen_items = [item for item in items if self.image_media_priority(item) in {0, 1}]
        # 历史工作流偶尔会保留已经失效的 cloud-tasks 地址，同时还会带有
        # 同一次 4B 推理实际归档到 dam 桶的候选帧。优先使用可读取的那一张，
        # 不能让失效地址把真正的 4B 抽帧排除在报告之外。
        ordered_items = qwen_items or items
        selected = dict(ordered_items[0])
        for item in ordered_items:
            for reference in self.image_item_references(item):
                if not self.read_minio_or_http_bytes(reference):
                    continue
                selected = dict(item)
                selected["url"] = reference
                break
            else:
                continue
            break
        selected["role"] = "model_representative"
        selected.setdefault("source", "model_representative")
        selected.setdefault("caption", "4B 代表性抽帧")
        return selected

    @staticmethod
    def image_item_references(item: dict[str, Any]) -> list[str]:
        """Return local persisted media before any legacy/cloud display URL."""
        references: list[str] = []
        source = item.get("source")
        if isinstance(source, dict):
            bucket = str(source.get("bucket") or "").strip()
            object_name = str(
                source.get("object_name") or source.get("object_key") or source.get("path") or ""
            ).strip().lstrip("/")
            if bucket and object_name:
                references.append(f"{bucket}/{object_name}")
        references.append(str(item.get("url") or "").strip())
        return list(dict.fromkeys(reference for reference in references if reference))

    def is_readable_image_item(self, item: Optional[dict[str, Any]]) -> bool:
        return bool(item and any(
            self.read_minio_or_http_bytes(reference)
            for reference in self.image_item_references(item)
        ))

    @staticmethod
    def select_linkage_report_images(
        evidence: list[SafetyEventEvidence],
    ) -> list[dict[str, Any]]:
        """Pick one representative image for each executed linkage object.

        Device actions retain every returned image in ``safety_event_evidence``.
        A report shows one image for each of drone, robot dog and staff when
        present, while additional images from the same action stay archived.
        """
        linkage_types = {"DRONE_IMAGE", "ROBOT_IMAGE", "STAFF_IMAGE"}
        linkage_sources = {"DRONE", "UAV", "ROBOT_DOG", "ROBOT", "STAFF"}
        linkage_labels = {
            "DRONE": "无人机",
            "UAV": "无人机",
            "DRONE_IMAGE": "无人机",
            "ROBOT_DOG": "机器狗",
            "ROBOT": "机器狗",
            "ROBOT_IMAGE": "机器狗",
            "STAFF": "人工处置",
            "STAFF_IMAGE": "人工处置",
        }
        selected: list[dict[str, Any]] = []
        selected_labels: set[str] = set()
        for row in evidence:
            if (
                str(row.evidence_type or "").upper() not in linkage_types
                and str(row.source_type or "").upper() not in linkage_sources
            ):
                continue
            if not row.file_url:
                continue
            linkage_label = (
                linkage_labels.get(str(row.source_type or "").upper())
                or linkage_labels.get(str(row.evidence_type or "").upper())
                or "联动设备"
            )
            if linkage_label in selected_labels:
                continue
            description = row.description or "现场取证"
            selected.append({
                "url": row.file_url,
                "caption": f"{linkage_label}联动代表性取证图：{description}",
                "source": "linkage_action_representative",
                "role": "linkage_representative",
                "linkage_label": linkage_label,
            })
            selected_labels.add(linkage_label)
        return selected

    @classmethod
    def select_linkage_report_image(
        cls,
        evidence: list[SafetyEventEvidence],
    ) -> Optional[dict[str, Any]]:
        """Compatibility helper for callers that need just the first object."""
        items = cls.select_linkage_report_images(evidence)
        return items[0] if items else None

    def extend_image_media_items(self, items: list[dict[str, Any]], value: Any) -> None:
        if not value:
            return
        if isinstance(value, list):
            for item in value:
                self.extend_image_media_items(items, item)
            return
        if isinstance(value, dict):
            media_type = str(value.get("type") or value.get("media_type") or "image").lower()
            if media_type and media_type not in {"image", "photo", "snapshot", "frame"}:
                return
            url = (
                value.get("url")
                or value.get("file_url")
                or value.get("image_url")
                or value.get("path")
                or value.get("annotated_ref")
                or value.get("object_url")
            )
            object_name = value.get("object_name") or value.get("object_key") or value.get("annotated_object_key")
            if not url and value.get("bucket") and object_name:
                url = f"{value.get('bucket')}/{object_name}"
            if self.is_yolo_detection_image(value, url):
                return
            if url and "{{" not in str(url) and "}}" not in str(url):
                items.append({
                    "url": str(url),
                    "caption": str(value.get("caption") or value.get("description") or ""),
                    "role": value.get("role"),
                    "source": value.get("source"),
                    "selected_by": value.get("selected_by"),
                    "timestamp_seconds": value.get("timestamp_seconds") or value.get("frame_time_sec"),
                })
            return
        if isinstance(value, str) and value and "{{" not in value and "}}" not in value:
            if self.is_yolo_detection_image({}, value):
                return
            items.append({"url": value, "caption": ""})

    @staticmethod
    def image_media_priority(item: dict[str, Any]) -> int:
        text = (
            f"{item.get('role') or ''} {item.get('source') or ''} "
            f"{item.get('selected_by') or ''} {item.get('caption') or ''} {item.get('url') or ''}"
        ).lower()
        if "qwen4b_selected_representative_frame" in text or "qwen4b_action_reasoning" in text:
            return 0
        if "qwen4b-proxy-media" in text or "qwen4b_review_frame_candidate" in text or "qwen4b_representative_frame_candidate" in text:
            return 1
        if "workflow-media" in text or "key_frame" in text:
            return 2
        if "qwen_screening" in text or "/camera/" in text:
            return 9
        return 5

    @staticmethod
    def is_yolo_detection_image(value: dict[str, Any], url: Any = None) -> bool:
        text = (
            f"{value.get('role') or ''} {value.get('source') or ''} "
            f"{value.get('object_name') or ''} {value.get('object_key') or ''} {url or ''}"
        ).lower()
        if "qwen4b-proxy-media" in text:
            return False
        return "annotated_detection_frame" in text or "workflow/yolo-detections" in text

    def collect_video_items(
        self,
        workflow_payload: dict[str, Any],
        visual: dict[str, Any],
        evidence: list[SafetyEventEvidence],
    ) -> list[dict[str, Any]]:
        items: list[dict[str, Any]] = []
        screening = visual.get("screening") if isinstance(visual.get("screening"), dict) else {}
        self.extend_media_items(items, visual.get("video_urls"))
        self.extend_media_items(items, visual.get("videos"))
        self.extend_media_items(items, screening.get("video_urls"))
        self.extend_media_items(items, screening.get("videos"))
        self.extend_media_items(items, self.find_nested_values(workflow_payload, "videos"))
        self.extend_media_items(items, self.find_nested_values(workflow_payload, "video_urls"))
        self.extend_video_media_objects(items, self.find_nested_values(workflow_payload, "media_objects"))
        self.extend_video_media_objects(items, self.find_nested_values(workflow_payload, "cloud_media_objects"))
        for row in evidence:
            if str(row.evidence_type or "").upper() == "VIDEO":
                items.append({"url": row.file_url, "caption": row.description or "事件证据视频"})
        return self.unique_media_items(items)

    def extend_video_media_objects(self, items: list[dict[str, Any]], value: Any) -> None:
        if not value:
            return
        if isinstance(value, list):
            for item in value:
                self.extend_video_media_objects(items, item)
            return
        if not isinstance(value, dict):
            return
        media_type = str(value.get("type") or value.get("media_type") or "").lower()
        url = (
            value.get("url")
            or value.get("file_url")
            or value.get("video_url")
            or value.get("path")
            or value.get("object_url")
        )
        object_name = value.get("object_name") or value.get("object_key")
        if not url and value.get("bucket") and object_name:
            url = f"{value.get('bucket')}/{object_name}"
        text = f"{media_type} {url or ''} {object_name or ''}".lower()
        if "video" not in media_type and not text.endswith((".mp4", ".mov", ".m4v", ".webm")):
            return
        if not url or "{{" in str(url) or "}}" in str(url):
            return
        items.append({
            "url": str(url),
            "caption": str(value.get("caption") or value.get("description") or ""),
            "role": value.get("role"),
            "source": value.get("source"),
            "content_type": value.get("content_type"),
        })

    def extend_media_items(self, items: list[dict[str, Any]], value: Any) -> None:
        if not value:
            return
        if isinstance(value, str):
            if "{{" in value or "}}" in value:
                return
            items.append({"url": value, "caption": ""})
            return
        if isinstance(value, dict):
            url = (
                value.get("url")
                or value.get("file_url")
                or value.get("image_url")
                or value.get("video_url")
                or value.get("object_url")
            )
            object_name = value.get("object_name") or value.get("object_key")
            if not url and value.get("bucket") and object_name:
                url = f"{value.get('bucket')}/{object_name}"
            if url:
                if "{{" in str(url) or "}}" in str(url):
                    return
                items.append({
                    "url": str(url),
                    "caption": str(value.get("caption") or value.get("description") or ""),
                    "role": value.get("role"),
                    "source": value.get("source"),
                    "content_type": value.get("content_type"),
                })
            return
        if isinstance(value, list):
            for item in value:
                self.extend_media_items(items, item)

    def unique_media_items(self, items: list[dict[str, Any]]) -> list[dict[str, Any]]:
        seen: set[str] = set()
        result = []
        for item in items:
            url = str(item.get("url") or "").strip()
            if not url or url in seen:
                continue
            bucket, object_name = self.parse_minio_reference(url)
            dedupe_key = f"{bucket}/{object_name}" if bucket and object_name else url
            if dedupe_key in seen:
                continue
            seen.add(dedupe_key)
            result.append(item)
        return result

    def find_nested_values(self, value: Any, key: str) -> list[Any]:
        values: list[Any] = []
        if isinstance(value, dict):
            for current_key, current_value in value.items():
                if current_key == key:
                    values.append(current_value)
                values.extend(self.find_nested_values(current_value, key))
        elif isinstance(value, list):
            for item in value:
                values.extend(self.find_nested_values(item, key))
        return values

    def download_first_image(self, items: list[dict[str, Any]]) -> Optional[Path]:
        for item in items:
            content = self.read_minio_or_http_bytes(str(item.get("url") or ""))
            if not content:
                continue
            suffix = ".jpg"
            parsed_suffix = Path(urlparse(str(item.get("url") or "")).path).suffix.lower()
            if parsed_suffix in {".jpg", ".jpeg", ".png", ".webp"}:
                suffix = parsed_suffix
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp.write(content)
                return Path(tmp.name)
        return None

    def select_evidence_image(
        self,
        image_items: list[dict[str, Any]],
        video_items: list[dict[str, Any]],
    ) -> Optional[Path]:
        model_item = next(
            (item for item in image_items if str(item.get("role") or "") == "model_representative"),
            None,
        )
        linkage_items = [
            item for item in image_items
            if str(item.get("role") or "") == "linkage_representative"
        ]
        report_items: list[tuple[dict[str, Any], str]] = []
        if model_item:
            report_items.append((model_item, "4B 代表性抽帧"))
        report_items.extend(
            (item, f"{item.get('linkage_label') or '联动设备'}联动取证")
            for item in linkage_items
        )
        if len(report_items) >= 2:
            composite = self.compose_evidence_images(report_items)
            if composite:
                return composite
        return self.download_first_image([
            item for item in (model_item, *linkage_items, *image_items) if item
        ])

    def compose_evidence_images(
        self,
        items: list[tuple[dict[str, Any], str]],
    ) -> Optional[Path]:
        """Compose representative frames for the single DOCX image slot."""
        panels: list[tuple[Image.Image, str]] = []
        for item, title in items:
            content = self.read_minio_or_http_bytes(str(item.get("url") or ""))
            if not content:
                continue
            try:
                with Image.open(io.BytesIO(content)) as source:
                    panels.append((ImageOps.exif_transpose(source).convert("RGB"), title))
            except Exception as exc:
                logger.debug("报告证据图合成读取失败 {}: {}", item.get("url"), exc)
        if len(panels) < 2:
            return None

        # 两张时沿用上下布局；三张及以上转为双列网格，既保留每个联动对象，
        # 也避免单张合成图超过一页可用高度。
        columns = 1 if len(panels) <= 2 else 2
        panel_width, panel_height = (1080, 500) if columns == 1 else (620, 330)
        title_height, gap, margin = (46, 26, 12) if columns == 1 else (38, 18, 12)
        panel_block_height = title_height + panel_height
        rows = (len(panels) + columns - 1) // columns
        canvas = Image.new(
            "RGB",
            (
                panel_width * columns + gap * (columns - 1) + margin * 2,
                panel_block_height * rows + gap * (rows - 1) + margin * 2,
            ),
            "white",
        )
        drawer = ImageDraw.Draw(canvas)
        try:
            title_font = ImageFont.truetype(
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                28 if columns == 1 else 20,
            )
        except OSError:
            title_font = ImageFont.load_default()
        for index, (image, title) in enumerate(panels):
            row, column = divmod(index, columns)
            block_x = margin + column * (panel_width + gap)
            block_y = margin + row * (panel_block_height + gap)
            title_box = drawer.textbbox((0, 0), title, font=title_font)
            title_x = block_x + (panel_width - (title_box[2] - title_box[0])) // 2
            drawer.text((title_x, block_y + 5), title, fill="#1f2937", font=title_font)
            panel = ImageOps.contain(image, (panel_width, panel_height))
            offset_x = block_x + (panel_width - panel.width) // 2
            offset_y = block_y + title_height + (panel_height - panel.height) // 2
            canvas.paste(panel, (offset_x, offset_y))
        with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
            canvas.save(tmp, format="JPEG", quality=92)
            return Path(tmp.name)

    def extract_frame_items_from_videos(
        self,
        video_items: list[dict[str, Any]],
        *,
        event_name: str,
        analysis_text: str,
        workflow_insight: dict[str, Any],
    ) -> list[dict[str, Any]]:
        """Legacy fallback for reports generated before the workflow returned representative frames."""
        items: list[dict[str, Any]] = []
        for index, item in enumerate(video_items[:2], 1):
            frame_path = self.extract_video_frame(str(item.get("url") or ""))
            if frame_path:
                items.append({
                    "url": str(frame_path),
                    "caption": item.get("caption") or f"事件证据视频兜底抽帧{index}",
                    "source": "legacy_video_frame_fallback",
                })
        return items

    def extract_video_frame(self, value: str) -> Optional[Path]:
        candidates = self.extract_video_candidate_frames(value)
        if not candidates:
            return None
        middle = candidates[len(candidates) // 2]
        selected_path = Path(str(middle["path"]))
        for candidate in candidates:
            path = Path(str(candidate.get("path") or ""))
            if path != selected_path:
                try:
                    path.unlink(missing_ok=True)
                except OSError:
                    pass
        return selected_path if selected_path.exists() else None

    def extract_video_candidate_frames(self, value: str, count: int = 4) -> list[dict[str, Any]]:
        video_bytes = self.read_minio_or_http_bytes(
            value,
            allowed_suffixes={".mp4", ".mov", ".m4v", ".webm"},
            allowed_content_prefixes={"video/"},
        )
        if not video_bytes:
            return []
        ffmpeg = shutil.which(settings.FFMPEG_BIN) or shutil.which("ffmpeg")
        if not ffmpeg:
            logger.warning("未找到 FFmpeg，无法从事件视频抽帧")
            return []
        input_path: Optional[Path] = None
        output_paths: list[Path] = []
        try:
            with tempfile.NamedTemporaryFile(suffix=".mp4", delete=False) as video_tmp:
                video_tmp.write(video_bytes)
                input_path = Path(video_tmp.name)
            duration = self.probe_video_duration(input_path)
            timestamps = self.candidate_frame_timestamps(duration, count=count)
            candidates: list[dict[str, Any]] = []
            for frame_index, timestamp in enumerate(timestamps, 1):
                with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as frame_tmp:
                    output_path = Path(frame_tmp.name)
                output_paths.append(output_path)
                command = [
                    ffmpeg,
                    "-hide_banner",
                    "-loglevel",
                    "error",
                    "-ss",
                    f"{timestamp:.3f}",
                    "-i",
                    str(input_path),
                    "-frames:v",
                    "1",
                    "-vf",
                    "scale='if(gte(iw,ih),min(iw,1280),-2)':'if(gte(iw,ih),-2,min(ih,720))'",
                    "-q:v",
                    "4",
                    "-y",
                    str(output_path),
                ]
                result = subprocess.run(
                    command,
                    stdin=subprocess.DEVNULL,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=12,
                    check=False,
                )
                if result.returncode == 0 and output_path.exists() and output_path.stat().st_size > 0:
                    candidates.append({
                        "index": frame_index,
                        "timestamp": round(timestamp, 3),
                        "path": output_path,
                    })
                else:
                    message = result.stderr.decode("utf-8", errors="replace").strip()
                    logger.debug("候选代表帧抽取失败 {} @{}s: {}", value, timestamp, message[:200])
            return candidates
        except Exception as exc:
            logger.warning("事件视频抽帧异常 {}: {}", value, exc)
        finally:
            if input_path:
                try:
                    input_path.unlink(missing_ok=True)
                except OSError:
                    pass
        for output_path in output_paths:
            try:
                output_path.unlink(missing_ok=True)
            except OSError:
                pass
        return []

    @staticmethod
    def probe_video_duration(input_path: Path) -> float:
        ffprobe = shutil.which("ffprobe")
        if not ffprobe:
            return 0.0
        try:
            result = subprocess.run(
                [
                    ffprobe,
                    "-v",
                    "error",
                    "-show_entries",
                    "format=duration",
                    "-of",
                    "default=noprint_wrappers=1:nokey=1",
                    str(input_path),
                ],
                stdin=subprocess.DEVNULL,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=5,
                check=False,
            )
            return max(0.0, float((result.stdout or b"").decode().strip() or 0.0))
        except Exception:
            return 0.0

    @staticmethod
    def candidate_frame_timestamps(duration: float, count: int = 4) -> list[float]:
        count = max(1, min(int(count or 4), 6))
        if duration > 1.2:
            start = min(0.6, duration * 0.12)
            end = max(start, duration - min(0.4, duration * 0.08))
            if count == 1:
                return [duration / 2]
            return [
                start + (end - start) * index / (count - 1)
                for index in range(count)
            ]
        return [0.2 + index * 0.8 for index in range(count)]

    def read_minio_or_http_bytes(
        self,
        value: str,
        *,
        allowed_suffixes: Optional[set[str]] = None,
        allowed_content_prefixes: Optional[set[str]] = None,
    ) -> Optional[bytes]:
        allowed_suffixes = allowed_suffixes or {".jpg", ".jpeg", ".png", ".webp"}
        allowed_content_prefixes = allowed_content_prefixes or {"image/"}
        local_path = Path(str(value or ""))
        if local_path.is_absolute() and local_path.exists():
            suffix = local_path.suffix.lower()
            if suffix in allowed_suffixes:
                try:
                    return local_path.read_bytes()
                except OSError as exc:
                    logger.debug("读取本地图像失败 {}: {}", value, exc)
        bucket, object_name = self.parse_minio_reference(value)
        if bucket and object_name and not minio_service.client:
            try:
                minio_service.connect()
            except Exception as exc:
                logger.debug("MinIO 懒连接失败: {}", exc)
        if bucket and object_name and minio_service.client:
            try:
                response = minio_service.client.get_object(bucket, object_name)
                try:
                    return response.read()
                finally:
                    response.close()
                    response.release_conn()
            except Exception as exc:
                logger.debug("读取 MinIO 图像失败 {}: {}", value, exc)
        parsed = urlparse(value)
        if parsed.scheme in {"http", "https"}:
            urls = [value]
            if parsed.hostname in {"localhost", "127.0.0.1"} and parsed.port == 9000:
                endpoint = getattr(settings, "QWEN_CAMERA_SCREENING_MINIO_ENDPOINT", "") or "172.17.0.1:9000"
                urls.append(urlunparse(parsed._replace(netloc=endpoint)))
            for url in urls:
                try:
                    with urllib.request.urlopen(url, timeout=5) as response:
                        content_type = response.headers.get("Content-Type") or ""
                        data = response.read()
                    if data and (
                        any(content_type.startswith(prefix) for prefix in allowed_content_prefixes)
                        or Path(parsed.path).suffix.lower() in allowed_suffixes
                    ):
                        return data
                except Exception as exc:
                    logger.debug("HTTP读取图像失败 {}: {}", url, exc)
        return None

    def parse_minio_reference(self, value: str) -> tuple[Optional[str], Optional[str]]:
        text = str(value or "").strip()
        if not text:
            return None, None
        parsed = urlparse(text)
        if parsed.scheme in {"http", "https"}:
            parts = [unquote(part) for part in parsed.path.strip("/").split("/") if part]
            if len(parts) >= 2:
                return parts[0], "/".join(parts[1:])
            return None, None
        if "/" in text and not text.startswith("/"):
            bucket, object_name = text.split("/", 1)
            if bucket and object_name:
                return bucket, object_name
        return None, None

    def source_label(self, instance: SafetyEventInstance, source: Optional[DataSource], selected: dict[str, Any]) -> str:
        return f"{self.source_trigger_label(instance, source)} · {selected['source_label']}"

    @staticmethod
    def source_trigger_label(instance: SafetyEventInstance, source: Optional[DataSource]) -> str:
        source_type = str(instance.source_type or getattr(source, "source_type", "") or "").lower()
        return {"camera": "摄像头触发", "sensor": "传感器触发"}.get(source_type, source_type or "事件触发")

    def location(self, source: Optional[DataSource], camera: Optional[Camera], visual: dict[str, Any]) -> str:
        values = [
            getattr(camera, "install_address", None),
            visual.get("zone_name"),
            visual.get("camera_name"),
            getattr(source, "source_name", None),
        ]
        return self.unique_location_text(values)

    @staticmethod
    def unique_location_text(values: list[Any]) -> str:
        cleaned: list[str] = []
        for value in values:
            text = re.sub(r"\s+", "", str(value or "").strip())
            if not text or text == "—":
                continue
            if any(text == existing or text in existing or existing in text for existing in cleaned):
                if any(existing in text and existing != text for existing in cleaned):
                    cleaned = [existing for existing in cleaned if existing not in text]
                    cleaned.append(text)
                continue
            cleaned.append(text)
        return " · ".join(cleaned) or "—"

    def result_label(self, instance: SafetyEventInstance) -> str:
        if instance.state == "RESOLVED":
            if instance.status == "FALSE_ALARM":
                return "误报关闭"
            return "已闭环"
        status = str(instance.status or "").upper()
        if status in {"PENDING", "PROCESSING"}:
            return "报告已生成"
        return STATUS_NAMES.get(status, "报告已生成")

    def completed_at(
        self,
        instance: SafetyEventInstance,
        timeline: Optional[list[SafetyEventTimelineLog]] = None,
    ) -> str:
        end = self.completion_time(instance, timeline)
        return self.format_datetime(end) if end else "—"

    def event_duration(
        self,
        instance: SafetyEventInstance,
        timeline: Optional[list[SafetyEventTimelineLog]] = None,
    ) -> str:
        start = instance.started_at
        end = self.completion_time(instance, timeline) or instance.last_observed_at or dt.datetime.now()
        if not start or not end:
            return "—"
        seconds = max(0, int((end - start).total_seconds()))
        minutes, sec = divmod(seconds, 60)
        hours, minutes = divmod(minutes, 60)
        if hours:
            return f"{hours}小时{minutes}分钟{sec}秒"
        if minutes:
            return f"{minutes}分钟{sec}秒"
        return f"{sec}秒"

    @staticmethod
    def completion_time(
        instance: SafetyEventInstance,
        timeline: Optional[list[SafetyEventTimelineLog]] = None,
    ) -> Optional[dt.datetime]:
        if timeline:
            candidates = [
                row.create_time
                for row in timeline
                if row.create_time
                and str(row.log_type or "").upper() in {"ACTION", "REPORT", "DAM_WORKFLOW"}
            ]
            if candidates:
                return max(candidates)
        return instance.resolved_at

    def instance_no_parts(self, instance_no: str) -> tuple[str, str]:
        text = str(instance_no or "").strip()
        if "_" not in text:
            return text, ""
        parts = text.split("_")
        if len(parts) >= 3:
            return "_".join(parts[:2]), "_".join(parts[2:])
        return text, ""

    def emergency_level(self, instance: SafetyEventInstance) -> str:
        risk = str(instance.max_risk_level or instance.risk_level or "").upper()
        return {"HIGH": "I级关注", "MEDIUM": "II级关注", "LOW": "III级关注"}.get(risk, "待确认")

    def confidence_label(self, selected: dict[str, Any]) -> str:
        confidence = self.numeric_value(self.find_in_selected(selected, "confidence"))
        if confidence is None:
            return "—"
        return f"{confidence * 100:.1f}%"

    def event_summary(
        self,
        instance: SafetyEventInstance,
        event: EventLibrary,
        source: Optional[DataSource],
        visual: dict[str, Any],
        insight: dict[str, Any],
    ) -> str:
        camera_name = visual.get("camera_name") or getattr(source, "source_name", None) or "现场摄像头"
        event_name = getattr(event, "event_name", None) or instance.summary or "安全事件"
        suspect_note = self.screening_suspect_note(visual, event_name=event_name)
        if suspect_note:
            return self.compact(f"{camera_name}触发{event_name}，{suspect_note}", 180)
        qwen_summary = insight.get("qwen_summary")
        if self.looks_like_model_thinking(str(qwen_summary or "")):
            qwen_summary = ""
        qwen4b_summary = self.clean_model_output_field(insight.get("qwen4b_conclusion"), allow_empty=True)
        if qwen4b_summary and qwen4b_summary != "—":
            return f"{camera_name}触发{event_name}，智能分析结论：{qwen4b_summary}"
        if qwen_summary:
            return self.sentence_safe_limit(f"{camera_name}触发{event_name}，初筛摘要：{qwen_summary}。", 520)
        detected = insight.get("specialized_class_label")
        if detected:
            return f"{camera_name}触发{event_name}，专有模型复核结果为{detected}。"
        return self.sentence_safe_limit(instance.summary or event_name, 520)

    def key_observation(
        self,
        workflow_payload: dict[str, Any],
        visual: dict[str, Any],
        selected_text: str,
        insight: Optional[dict[str, Any]] = None,
        selected: Optional[dict[str, Any]] = None,
    ) -> str:
        insight = insight or {}
        parts = []
        suspect_note = self.screening_suspect_note(visual)
        if suspect_note:
            parts.append(suspect_note)
        screening = visual.get("screening") if isinstance(visual.get("screening"), dict) else {}
        if screening.get("summary") and not self.looks_like_model_thinking(str(screening.get("summary"))):
            parts.append(str(screening.get("summary")))
        if not parts:
            selected_observation = self.find_in_selected(selected, "key_observation") if selected else None
            if not selected_observation and selected and selected.get("source") == "qwen35b":
                selected_observation = (
                    self.find_in_selected(selected, "detailed_scene_analysis")
                    or self.find_in_selected(selected, "scene_analysis")
                )
            if selected_observation:
                parts.append(self.compact(str(selected_observation), 220))
            elif insight.get("qwen4b_risk_reasoning"):
                parts.append(self.compact(str(insight.get("qwen4b_risk_reasoning")), 220))
        main_class = insight.get("specialized_class")
        confidence = insight.get("specialized_confidence")
        if main_class:
            parts.append(f"专有模型复核：{insight.get('specialized_class_label') or main_class}")
        if confidence is not None:
            try:
                parts.append(f"置信度：{float(confidence) * 100:.1f}%")
            except (TypeError, ValueError):
                parts.append(f"置信度：{confidence}")
        if parts:
            return "；".join(parts[:3])
        return self.compact(selected_text, 180)

    def report_risk_label(self, instance: SafetyEventInstance, selected: dict[str, Any]) -> str:
        """Use the selected final model result consistently in the report header."""
        current = str(instance.max_risk_level or instance.risk_level or "").upper()
        if current in RISK_NAMES:
            return RISK_NAMES[current]
        value = self.find_in_selected(selected, "risk_level")
        text = str(value or "").strip().lower()
        labels = {
            "high": "高风险",
            "critical": "严重风险",
            "medium": "中风险",
            "low": "低风险",
            "高风险": "高风险",
            "严重风险": "严重风险",
            "中风险": "中风险",
            "低风险": "低风险",
        }
        if text in labels:
            return labels[text]
        return "低风险"

    def sync_instance_risk_from_report(
        self,
        instance: SafetyEventInstance,
        selected: dict[str, Any],
    ) -> None:
        """Promote the event risk when the final report finds a higher risk.

        The trigger stores the initial screening risk. A completed 4B/35B
        report is the final review result, so its higher level must be visible
        in the event center as well as in the DOCX. Never downgrade an event's
        recorded maximum risk during report generation.
        """
        value = self.find_in_selected(selected, "risk_level")
        text = str(value or "").strip().lower()
        final_risk = {
            "critical": "HIGH",
            "严重风险": "HIGH",
            "high": "HIGH",
            "高风险": "HIGH",
            "medium": "MEDIUM",
            "中风险": "MEDIUM",
            "low": "LOW",
            "低风险": "LOW",
        }.get(text)
        if not final_risk:
            return
        rank = {"LOW": 1, "MEDIUM": 2, "HIGH": 3}
        current_max = str(instance.max_risk_level or instance.risk_level or "LOW").upper()
        if rank.get(final_risk, 0) > rank.get(current_max, 0):
            instance.max_risk_level = final_risk
        if rank.get(final_risk, 0) >= rank.get(str(instance.risk_level or "LOW").upper(), 0):
            instance.risk_level = final_risk

    def first_nested_value(self, value: Any, key: str) -> Any:
        values = self.find_nested_values(value, key)
        for item in values:
            if item in (None, "", []):
                continue
            if isinstance(item, dict):
                continue
            if isinstance(item, list):
                continue
            else:
                return item
        return None

    def source_summary(self, image_items: list[dict[str, Any]], video_items: list[dict[str, Any]], selected: dict[str, Any]) -> str:
        pieces = [f"采用{selected['source_label']}生成分析结论"]
        if video_items:
            pieces.append(f"关联事件证据视频 {len(video_items)} 段")
        if image_items:
            model_count = self.model_image_count(image_items)
            linkage_labels = self.linkage_image_labels(image_items)
            if model_count:
                pieces.append(f"关联 4B 代表性抽帧 {model_count} 张")
            if linkage_labels:
                pieces.append(f"关联{self.format_linkage_image_labels(linkage_labels)}联动代表性取证图 {sum(linkage_labels.values())} 张")
            elif not model_count:
                pieces.append(f"关联抽帧图像 {len(image_items)} 张")
        return "；".join(pieces) + "。"

    def evidence_summary(self, image_items: list[dict[str, Any]], video_items: list[dict[str, Any]]) -> str:
        if not image_items and not video_items:
            return "本次事件未记录可用图像或视频证据。"
        pieces = []
        if video_items:
            pieces.append(f"已归档 {len(video_items)} 段事件证据视频")
        if image_items:
            model_count = self.model_image_count(image_items)
            linkage_labels = self.linkage_image_labels(image_items)
            if model_count:
                pieces.append(f"已归档 {model_count} 张 4B 代表性抽帧")
            if linkage_labels:
                pieces.append(f"已归档 {self.format_linkage_image_labels(linkage_labels)}联动代表性取证图 {sum(linkage_labels.values())} 张")
            elif not model_count:
                pieces.append(f"已归档 {len(image_items)} 张关键帧/检测图像")
        return "；".join(pieces) + "，用于支撑本次事件研判。"

    def evidence_caption(self, image_items: list[dict[str, Any]], video_items: list[dict[str, Any]]) -> str:
        return ""

    def trigger_summary(
        self,
        instance: SafetyEventInstance,
        event: EventLibrary,
        visual: dict[str, Any],
        insight: dict[str, Any],
    ) -> str:
        event_name = getattr(event, "event_name", None) or instance.summary or "安全事件"
        camera_name = visual.get("camera_name") or "现场摄像头"
        suspect_note = self.screening_suspect_note(visual)
        if suspect_note:
            return f"{camera_name}在{self.format_datetime(instance.started_at)}触发{event_name}，{suspect_note}"
        summary = insight.get("qwen_summary") or "触发时未记录初筛摘要"
        return f"{camera_name}在{self.format_datetime(instance.started_at)}触发{event_name}，初筛摘要为：{summary}。"

    def screening_suspect_note(self, visual: dict[str, Any], event_name: str = "") -> str:
        screening = visual.get("screening") if isinstance(visual.get("screening"), dict) else {}
        if not screening:
            return ""
        pieces = []
        boat_confidence = self.numeric_value(screening.get("boat_confidence"))
        person_confidence = self.numeric_value(screening.get("person_confidence"))
        event_text = str(event_name or "")
        person_event = any(token in event_text for token in ("人员", "亲水", "涉水", "滩涂", "入侵"))
        boat_event = any(token in event_text for token in ("船", "捕鱼", "电鱼", "偷捕"))
        if boat_event and self.truthy(screening.get("possible_boat")) and not self.truthy(screening.get("boat_present")):
            confidence = f"{boat_confidence * 100:.1f}%" if boat_confidence is not None else "低置信"
            pieces.append(f"初筛标记疑似船只/疑似捕鱼，置信度{confidence}，未达到明确确认阈值，需结合后续模型和现场证据复核")
        if (person_event or not boat_event) and self.truthy(screening.get("possible_person")) and not self.truthy(screening.get("person_present")):
            confidence = f"{person_confidence * 100:.1f}%" if person_confidence is not None else "低置信"
            pieces.append(f"初筛标记疑似人员，置信度{confidence}，未达到明确确认阈值，需结合后续模型和现场证据复核")
        return "；".join(pieces) + ("。" if pieces else "")

    def screening_summary(self, visual: dict[str, Any]) -> str:
        screening = visual.get("screening") if isinstance(visual.get("screening"), dict) else {}
        if not screening:
            return "未获取到摄像头初筛结构化结果。"
        labels = [
            ("洪水", "flood_detected", "flood_confidence"),
            ("泥石流", "mudslide_detected", "mudslide_confidence"),
            ("滑坡", "landslide_detected", "landslide_confidence"),
            ("地震", "earthquake_detected", "earthquake_confidence"),
            ("人员", "person_present", "person_confidence"),
            ("船只/捕鱼", "boat_present", "boat_confidence"),
        ]
        hits = []
        negatives = []
        for label, flag_key, confidence_key in labels:
            flag = screening.get(flag_key)
            confidence = self.numeric_value(screening.get(confidence_key))
            text = f"{label}({confidence * 100:.1f}%)" if confidence is not None else label
            if str(flag) in {"1", "True", "true"} or flag == 1 or flag is True:
                hits.append(text)
            else:
                negatives.append(label)
        hit_text = "、".join(hits) if hits else "未命中明确场景"
        negative_text = "、".join(negatives[:6]) if negatives else "无"
        possible_hits = []
        for label, flag_key in (("疑似人员", "possible_person"), ("疑似船只", "possible_boat")):
            flag = screening.get(flag_key)
            if str(flag) in {"1", "True", "true"} or flag == 1 or flag is True:
                possible_hits.append(label)
        possible_text = "、".join(possible_hits) if possible_hits else "无"
        return (f"初筛命中：{hit_text}；疑似待复核：{possible_text}；"
                f"未命中/排除：{negative_text}；风险等级：{screening.get('qwen_risk_level') or '—'}。")

    @staticmethod
    def truthy(value: Any) -> bool:
        return value is True or value == 1 or str(value).lower() in {"1", "true", "yes", "y"}

    def model_route_summary(self, workflow_payload: dict[str, Any], selected: dict[str, Any], event_name: str = "") -> str:
        execution = workflow_payload.get("execution_result") if isinstance(workflow_payload, dict) else {}
        nodes = (execution or {}).get("node_results") or []
        success_nodes = {
            str(row.get("node_id") or "")
            for row in nodes
            if isinstance(row, dict) and str(row.get("status") or "").lower() == "success"
        }
        route_label = self.route_label_for_event(event_name, workflow_payload)
        route_parts = []
        if "action_classify" in success_nodes:
            if route_label == "极端天气环境风险分析链路":
                route_parts.append("多源环境数据完成风险复核")
            elif route_label == "人员异常行为分析链路":
                route_parts.append("专有模型完成人员/目标复核")
            else:
                route_parts.append("专有模型完成灾害类别复核")
        if "action_reasoning" in success_nodes:
            route_parts.append("4B本地模型完成现场语义理解")
        if "action_report" in success_nodes:
            route_parts.append("35B云端模型完成增强研判与报告校核")
        route_text = "，".join(route_parts) if route_parts else "模型节点已完成可用结果回传"
        source_label = selected.get("source_label", "智能分析模型")
        return f"ECA触发后，智能路由进入{route_label}，{route_text}；本报告以{source_label}结果作为最终分析依据。"

    def route_label_for_event(self, event_name: str, workflow_payload: dict[str, Any]) -> str:
        text_parts = [event_name]
        if isinstance(workflow_payload, dict):
            for key in ("event_type", "event_name", "event_group"):
                if workflow_payload.get(key):
                    text_parts.append(str(workflow_payload.get(key)))
            plan = workflow_payload.get("workflow") or workflow_payload.get("dag") or workflow_payload.get("plan")
            if isinstance(plan, dict):
                for key in ("event_type", "event_name", "event_group", "description"):
                    if plan.get(key):
                        text_parts.append(str(plan.get(key)))
        text = " ".join(text_parts)
        if any(token in text for token in ("人员", "闯入", "亲水", "涉水", "船只", "偷捕", "电鱼", "person_behavior", "PERSON_SAFETY", "ILLEGAL_FISHING")):
            return "人员异常行为分析链路"
        if any(token in text for token in ("台风", "飓风", "大风", "强风", "烈风", "狂风", "暴风", "暴雨", "高温", "低温", "极高温", "极低温", "冰冻", "高湿", "低湿", "雨量", "环境", "extreme_weather")):
            return "极端天气环境风险分析链路"
        if any(token in text for token in ("泥石流", "滑坡", "洪水", "地震", "natural_disaster")):
            return "自然灾害分析链路"
        return "智能综合分析链路"

    def workflow_nodes_summary(self, workflow_payload: dict[str, Any]) -> str:
        execution = workflow_payload.get("execution_result") if isinstance(workflow_payload, dict) else {}
        rows = []
        label_map = {
            "start_0": "事件触发",
            "action_classify": "专有模型复核",
            "action_reasoning": "4B 场景理解",
            "action_report": "35B 增强分析",
            "end_0": "流程结束",
        }
        for row in (execution or {}).get("node_results") or []:
            if not isinstance(row, dict):
                continue
            node_id = str(row.get("node_id") or "")
            status = str(row.get("status") or "unknown")
            model_id = row.get("model_id")
            meta = row.get("request_meta") if isinstance(row.get("request_meta"), dict) else {}
            media = []
            if meta.get("video_count"):
                media.append(f"视频{meta.get('video_count')}段")
            if meta.get("image_count"):
                media.append(f"图片{meta.get('image_count')}张")
            media_text = f"，输入{'、'.join(media)}" if media else ""
            model_text = f"，模型ID {model_id}" if model_id else ""
            rows.append(f"{label_map.get(node_id, node_id)}：{status}{model_text}{media_text}")
        return "\n".join(rows) or "未记录模型节点执行明细。"

    def specialized_summary(self, insight: dict[str, Any]) -> str:
        label = insight.get("specialized_class_label") or "未获得专有模型类别"
        confidence = insight.get("specialized_confidence")
        confidence_text = f"{confidence * 100:.1f}%" if confidence is not None else "—"
        sampled = insight.get("sampled_frames") or "—"
        report = insight.get("classification_report")
        suffix = f"；模型说明：{self.compact(str(report), 220)}" if report else ""
        return f"复核类别：{label}；置信度：{confidence_text}；采样帧数：{sampled}{suffix}。"

    def node_analysis_summary(self, workflow_payload: dict[str, Any], node_id: str) -> str:
        inference = self.find_node_inference(workflow_payload, node_id)
        if not inference:
            return "该节点未返回可用分析。"
        source = inference.get("system_prompt_source") or inference.get("actor_name") or node_id
        risk = inference.get("risk_level") or self.find_in_value(inference, "risk_level") or "—"
        confidence = self.numeric_value(inference.get("confidence") or self.find_in_value(inference, "confidence"))
        confidence_text = f"{confidence * 100:.1f}%" if confidence is not None else "—"
        report = (
            inference.get("report")
            or self.find_in_value(inference, "scene_analysis")
            or self.find_in_value(inference, "detailed_scene_analysis")
            or "未生成文字摘要"
        )
        return f"来源：{source}；风险：{risk}；置信度：{confidence_text}；摘要：{self.compact(str(report), 420)}"

    def evidence_inventory(self, image_items: list[dict[str, Any]], video_items: list[dict[str, Any]]) -> str:
        parts = []
        if video_items:
            parts.append(f"事件证据视频 {len(video_items)} 段")
        if image_items:
            model_count = self.model_image_count(image_items)
            linkage_labels = self.linkage_image_labels(image_items)
            if model_count:
                parts.append(f"4B 代表性抽帧 {model_count} 张")
            if linkage_labels:
                parts.append(f"{self.format_linkage_image_labels(linkage_labels)}联动代表性取证图 {sum(linkage_labels.values())} 张")
            elif not model_count:
                parts.append(f"关键帧/检测图像 {len(image_items)} 张")
        return "，".join(parts) + "，原始文件已归档至事件证据库。" if parts else "未归档媒体证据。"

    def frame_evidence_summary(self, image_items: list[dict[str, Any]], video_items: list[dict[str, Any]]) -> str:
        model_frames = 0
        linkage_labels: dict[str, int] = {}
        for item in image_items:
            url = str(item.get("url") or "")
            role = str(item.get("role") or "")
            if role == "model_representative" or "workflow-media" in url or "qwen4b-proxy-media" in url or "yolo-detections" in url or "key_frame" in url:
                model_frames += 1
            elif role == "linkage_representative":
                label = str(item.get("linkage_label") or "联动设备").strip() or "联动设备"
                linkage_labels[label] = linkage_labels.get(label, 0) + 1
        parts = []
        if video_items:
            parts.append(f"事件证据视频{len(video_items)}段")
        if model_frames:
            parts.append(f"4B代表性抽帧{model_frames}张")
        if linkage_labels:
            parts.append(f"{self.format_linkage_image_labels(linkage_labels)}联动取证图片{sum(linkage_labels.values())}张")
        if not parts:
            return "未记录可用于报告展示的抽帧图片。"
        return "已归档" + "、".join(parts) + "；报告正文嵌入代表性画面，其余图片随事件证据一并留存。"

    @staticmethod
    def linkage_image_labels(image_items: list[dict[str, Any]]) -> dict[str, int]:
        labels: dict[str, int] = {}
        for item in image_items:
            if str(item.get("role") or "") != "linkage_representative":
                continue
            label = str(item.get("linkage_label") or "联动设备").strip() or "联动设备"
            labels[label] = labels.get(label, 0) + 1
        return labels

    @staticmethod
    def model_image_count(image_items: list[dict[str, Any]]) -> int:
        return sum(
            1
            for item in image_items
            if str(item.get("role") or "") == "model_representative"
        )

    @staticmethod
    def format_linkage_image_labels(labels: dict[str, int]) -> str:
        return "、".join(labels.keys())

    def linkage_evidence_summary(self, evidence: list[SafetyEventEvidence]) -> str:
        labels = {
            "DRONE": "无人机",
            "UAV": "无人机",
            "DRONE_IMAGE": "无人机",
            "DRONE_VIDEO": "无人机",
            "ROBOT_DOG": "机器狗",
            "ROBOT": "机器狗",
            "ROBOT_IMAGE": "机器狗",
            "ROBOT_VIDEO": "机器狗",
            "STAFF": "人工处置",
            "STAFF_IMAGE": "人工处置",
        }
        counts: dict[str, int] = {}
        for row in evidence:
            source_type = str(row.source_type or "").upper()
            evidence_type = str(row.evidence_type or "").upper()
            label = labels.get(source_type) or labels.get(evidence_type)
            if not label:
                continue
            counts[label] = counts.get(label, 0) + 1
        if counts:
            return "；".join(f"{label}联动取证{count}条" for label, count in counts.items()) + "，作为现场补充证据。"
        return "本次事件暂未记录机器狗或无人机联动取证；如后续联动设备上传图片/视频，将作为补充证据归档。"

    def analysis_limitations(
        self,
        selected: dict[str, Any],
        insight: dict[str, Any],
        image_items: list[dict[str, Any]],
        video_items: list[dict[str, Any]],
    ) -> str:
        limitations = self.find_in_selected(selected, "uncertainties")
        if isinstance(limitations, list) and limitations:
            return "；".join(str(item) for item in limitations if item)
        if isinstance(limitations, str) and limitations.strip():
            return limitations.strip()
        parts = []
        if not video_items:
            parts.append("缺少可回放事件视频")
        if not image_items:
            parts.append("报告未嵌入可用关键帧")
        if not insight.get("specialized_confidence"):
            parts.append("专有模型置信度未记录")
        return "；".join(parts) if parts else "未发现明显数据缺口，仍建议结合现场人工复核。"

    def follow_up_actions(self, instance: SafetyEventInstance, selected: dict[str, Any]) -> str:
        risk = str(instance.max_risk_level or instance.risk_level or "").upper()
        base = [
            "保留事件视频、关键帧、模型结果和人工处置记录。",
            "将事件结论同步至值班台账，复盘模型输出质量。",
        ]
        if risk == "HIGH":
            base.insert(0, "按高风险事件持续跟踪，闭环后安排现场或远程复核。")
        if selected.get("cloud_error"):
            base.append("云端增强恢复后，可按需重新生成增强报告。")
        return "\n".join(f"{idx}. {text}" for idx, text in enumerate(base, 1))

    def recommendations_text(self, selected: dict[str, Any]) -> str:
        value = self.find_in_selected(selected, "recommendations")
        if isinstance(value, list) and value:
            return "\n".join(f"{idx}. {item}" for idx, item in enumerate(value, 1) if item)
        if isinstance(value, str) and value.strip():
            return value.strip()
        return "1. 继续监测事件区域。\n2. 结合现场条件执行人工复核。\n3. 视风险变化升级联动处置。"

    @staticmethod
    def risk_escalation_summary(risk_escalation: dict[str, Any], supplemental_context: dict[str, Any]) -> str:
        reason = str(risk_escalation.get("reason") or "").strip()
        label = str(supplemental_context.get("label") or "").strip()
        if reason:
            return reason
        if label:
            return f"补充运行状态显示{label}，知识库风险复核提示需按更高风险等级处置。"
        return "补充运行状态和知识库检索结果提示本事件风险等级已调整。"

    @staticmethod
    def risk_escalation_knowledge_summary(risk_escalation: dict[str, Any]) -> str:
        hits = risk_escalation.get("knowledge_hits")
        if not isinstance(hits, list) or not hits:
            return ""
        lines = []
        for index, item in enumerate(hits[:3], start=1):
            if not isinstance(item, dict):
                continue
            title = item.get("document_title") or "知识库文档"
            clause = item.get("clause_id") or item.get("evidence_id") or ""
            lines.append(f"[{index}] 《{title}》" + (f"，条款 {clause}" if clause else ""))
        return "\n".join(lines)

    @staticmethod
    def risk_escalation_report_overrides(summary: str, supplemental_context: dict[str, Any]) -> dict[str, str]:
        label = str(supplemental_context.get("label") or "特殊工况").strip()
        condition = f"{label}条件下" if label else "特殊工况条件下"
        impact = (
            f"{condition}，滩涂及近水区域人员活动可能受到泄洪水流、水位上涨、岸坡湿滑和退避通道受限等因素影响，"
            "存在被困、冲刷、滑跌和溺水风险；即使单帧画面未显示直接遇险，也应按高风险事件纳入闭环处置。"
        )
        response = (
            "立即通过广播、云台跟踪和现场巡查联动开展劝离，通知值班人员核实人员位置和撤离通道；"
            "在人员撤离并经复核确认前，维持高风险等级和事件取证，必要时联动安保或属地应急力量到场处置。"
        )
        monitoring = (
            "持续跟踪人员是否仍停留在滩涂、近水边坡、泄洪影响区或下游管控区，重点复核水位变化、闸门状态、"
            "泄洪流量和退避路线；人员清离后保留带框检测帧、视频片段和处置记录作为闭环依据。"
        )
        conclusion = (
            f"{summary} 系统已按高风险事件生成报告并进入闭环处置；最终处置以人员撤离、现场复核和证据归档完成为准。"
        )
        return {
            "impact_assessment": impact,
            "response_plan": response,
            "monitoring_suggestions": monitoring,
            "conclusion": conclusion,
        }

    @staticmethod
    def merge_knowledge_summaries(*summaries: Any) -> str:
        lines: list[str] = []
        seen: set[str] = set()
        for summary in summaries:
            for raw in str(summary or "").splitlines():
                line = raw.strip()
                if not line:
                    continue
                content_key = re.sub(r"^\[\d+\]\s*", "", line)
                if content_key in seen:
                    continue
                seen.add(content_key)
                lines.append(line)
        return "\n".join(
            re.sub(r"^\[\d+\]", f"[{index}]", line)
            for index, line in enumerate(lines, start=1)
        )

    def final_report_field(self, selected: dict[str, Any], key: str, fallback: Any = "") -> str:
        citation_map = self.knowledge_citation_index_map(selected)
        value = self.find_in_selected(selected, key)
        if isinstance(value, list):
            text = "；".join(str(item) for item in value if item)
            return self.apply_knowledge_citation_indexes(self.clean_model_output_field(text), citation_map)
        if isinstance(value, dict):
            return self.apply_knowledge_citation_indexes(
                self.clean_model_output_field(self.format_structured_value(value)),
                citation_map,
            )
        text = str(value or fallback or "").strip()
        return self.apply_knowledge_citation_indexes(self.clean_model_output_field(text), citation_map)

    def knowledge_citation_index_map(self, selected: dict[str, Any]) -> dict[str, int]:
        sources = selected.get("citation_sources")
        if not isinstance(sources, list):
            sources = self.find_in_selected(selected, "knowledge_sources") or selected.get("knowledge_sources")
        if not isinstance(sources, list):
            return {}
        citation_ids = {
            str(value).strip()
            for value in selected.get("citation_ids") or []
            if str(value).strip()
        }
        if citation_ids:
            cited_sources = []
            for item in sources:
                if not isinstance(item, dict):
                    continue
                identifiers = {
                    str(item.get("evidence_id") or "").strip(),
                    str(item.get("clause_id") or "").strip(),
                }
                chunk_id = str(item.get("chunk_id") or "").strip()
                if chunk_id:
                    identifiers.add(f"K{chunk_id}")
                if identifiers & citation_ids:
                    cited_sources.append(item)
            sources = cited_sources
        mapping: dict[str, int] = {}
        seen: set[tuple[str, str]] = set()
        index = 0
        for item in sources:
            if not isinstance(item, dict):
                continue
            title = str(item.get("document_title") or item.get("filename") or "").strip()
            section = str(item.get("section_path") or item.get("section_title") or "").strip()
            if not title:
                continue
            key = (title, section)
            if key not in seen:
                seen.add(key)
                index += 1
            evidence_id = item.get("evidence_id") or (f"K{item.get('chunk_id')}" if item.get("chunk_id") else "")
            if evidence_id:
                mapping[str(evidence_id)] = index
            if item.get("chunk_id"):
                mapping[f"K{item.get('chunk_id')}"] = index
            clause_id = str(item.get("clause_id") or "").strip()
            if clause_id:
                mapping[clause_id] = index
        return mapping

    def apply_knowledge_citation_indexes(self, text: str, citation_map: dict[str, int]) -> str:
        if not text or not citation_map:
            return text
        result = str(text)
        for evidence_id, index in sorted(citation_map.items(), key=lambda item: len(item[0]), reverse=True):
            if not evidence_id:
                continue
            result = result.replace(f"[{evidence_id}]", f"[{index}]")
            result = re.sub(rf"(?<!\[){re.escape(evidence_id)}(?!\])", f"[{index}]", result)
        # Never turn an unknown model-generated K-id into an arbitrary valid
        # index. That would make a fabricated citation look grounded. Keep
        # the prose, but remove only the unsupported marker.
        result = re.sub(r"结合知识库依据\s*\[K\d+\]\s*中", "结合现场证据", result)
        result = re.sub(r"知识库依据\s*\[K\d+\]\s*中", "现有证据", result)
        result = re.sub(r"\[K\d+\]", "", result)
        result = re.sub(r"(?<!\[)K\d+(?!\])", "", result)
        result = re.sub(r"结合知识库(?!依据)", "结合知识库依据", result)
        return result

    def normalize_report_context_citations(self, context: dict[str, Any], selected: dict[str, Any]) -> None:
        citation_map = self.knowledge_citation_index_map(selected)
        if not citation_map:
            workflow_insight = context.get("workflow_insight") or {}
            if isinstance(workflow_insight, dict):
                citation_map = self.knowledge_citation_index_map({
                    "raw_output": workflow_insight,
                })
        knowledge_summary = (
            (context.get("workflow_insight") or {}).get("knowledge_sources_summary")
            or context.get("knowledge_sources_summary")
            or ""
        )
        if not citation_map and str(knowledge_summary).strip():
            citation_map = {"K0": 1}
        if not citation_map:
            return
        for key, value in list(context.items()):
            if key in {"workflow_insight", "evidence_image"}:
                continue
            if isinstance(value, str):
                context[key] = self.apply_knowledge_citation_indexes(value, citation_map)

    def find_in_selected(self, selected: dict[str, Any], key: str) -> Any:
        return self.find_in_value(selected.get("raw_output"), key)

    def find_in_value(self, value: Any, key: str) -> Any:
        if isinstance(value, dict):
            if key in value and value.get(key) not in (None, "", []):
                return value.get(key)
            for child_key in ("inference_result", "final_report", "scene_analysis", "template_data", "docx_context", "output"):
                found = self.find_in_value(value.get(child_key), key)
                if found not in (None, "", []):
                    return found
        elif isinstance(value, list):
            for item in value:
                found = self.find_in_value(item, key)
                if found not in (None, "", []):
                    return found
        return None

    def timeline_summary(self, timeline: list[SafetyEventTimelineLog]) -> str:
        if not timeline:
            return "暂无处置时间线记录。"
        rows = []
        for row in timeline[-8:]:
            time_text = self.to_local_datetime(row.create_time).strftime("%H:%M:%S") if row.create_time else "--:--:--"
            rows.append(f"{time_text} {row.log_type}/{row.status}：{row.message}")
        return "\n".join(rows)

    def build_conclusion(
        self,
        selected: dict[str, Any],
        workflow_insight: dict[str, Any],
        cloud_note: str,
    ) -> str:
        citation_map = self.knowledge_citation_index_map(selected)
        finish = lambda value: self.apply_knowledge_citation_indexes(  # noqa: E731
            self.clean_model_output_field(value),
            citation_map,
        )
        event_profile = self.event_profile(workflow_insight)
        model_conclusion = self.model_report_conclusion(selected, workflow_insight, event_profile)
        if model_conclusion:
            if cloud_note:
                return finish(f"{cloud_note}\n{model_conclusion}")
            return finish(model_conclusion)

        if workflow_insight.get("suspected"):
            event_name = event_profile["event_name"]
            person_event = event_profile["person_event"]
            boat_event = event_profile["boat_event"]
            parts = []
            if boat_event and self.truthy(workflow_insight.get("possible_boat")) and not self.truthy(workflow_insight.get("boat_present")):
                parts.append("疑似船只/疑似捕鱼")
            if (person_event or not boat_event) and self.truthy(workflow_insight.get("possible_person")) and not self.truthy(workflow_insight.get("person_present")):
                parts.append("疑似人员亲水/滩涂活动")
            target_text = "、".join(parts) or event_name or "疑似目标"
            if person_event and not boat_event:
                conclusion = (
                    f"本次事件按{target_text}进入复核，已完成智能分析、处置联动和证据归档。"
                    "当前画面属于低置信远距离小目标线索，尚不足以认定为明确涉水或违规行为；"
                    "后续应结合连续视频、现场巡查和专有检测模型结果确认人员位置、活动轨迹及最终处置等级。"
                )
            else:
                conclusion = (
                    f"本次事件为低置信{target_text}触发，已完成智能分析、处置联动和证据归档。"
                    "当前证据不足以直接认定为明确违规行为，也不应表述为已排除目标；"
                    "后续应结合连续视频、现场巡查和专有模型复核结果确认最终处置等级。"
                )
            if cloud_note:
                return finish(f"{cloud_note}\n{conclusion}")
            return finish(conclusion)

        conclusion = self.final_report_field(selected, "conclusion", "")
        if self.looks_like_model_thinking(conclusion):
            conclusion = "—"
        if conclusion == "—":
            impact = self.final_report_field(selected, "impact_assessment", "")
            monitoring = self.final_report_field(selected, "monitoring_suggestions", "")
            if self.looks_like_model_thinking(impact):
                impact = str(workflow_insight.get("qwen4b_impact_assessment") or "").strip() or "—"
            if self.looks_like_model_thinking(monitoring):
                monitoring = str(workflow_insight.get("qwen4b_monitoring_suggestions") or "").strip() or "—"
            impact = self.valid_report_text(impact)
            monitoring = self.valid_report_text(monitoring)
            if impact and monitoring:
                conclusion = f"{impact} 后续应{monitoring.lstrip('建议').lstrip('：:')}"
            elif impact:
                conclusion = impact
            else:
                event_text = (
                    workflow_insight.get("qwen4b_conclusion")
                    or workflow_insight.get("qwen_summary")
                    or ""
                )
                event_text = self.valid_report_text(event_text)
                if event_text:
                    conclusion = (
                        f"本次事件已完成智能分析、处置联动和证据归档。{self.compact(event_text, 180)}"
                        "后续应结合现场巡查与连续监测结果确认处置效果。"
                    )
                else:
                    conclusion = (
                        "本次事件已完成智能分析、处置联动和证据归档。"
                        "后续应结合现场巡查、传感器连续监测和现场处置记录确认处置效果。"
                    )
        if cloud_note:
            return finish(f"{cloud_note}\n{conclusion}")
        return finish(conclusion)

    def model_report_conclusion(
        self,
        selected: dict[str, Any],
        workflow_insight: dict[str, Any],
        event_profile: dict[str, Any],
    ) -> str:
        candidates = [
            self.find_in_selected(selected, "conclusion"),
            self.find_in_selected(selected, "handling_conclusion"),
            self.find_in_selected(selected, "final_conclusion"),
            workflow_insight.get("qwen4b_conclusion"),
        ]
        for candidate in candidates:
            cleaned = self.clean_model_output_field(candidate, allow_empty=True)
            if not self.valid_model_conclusion(cleaned, event_profile):
                continue
            return cleaned
        return ""

    def valid_model_conclusion(self, conclusion: str, event_profile: dict[str, Any]) -> bool:
        text = self.valid_report_text(conclusion)
        if not text:
            return False
        if self.looks_truncated(text):
            return False
        if event_profile.get("person_event"):
            conflict_terms = ("船只", "船舶", "小船", "捕鱼", "电鱼", "偷捕", "非法捕捞", "疑似船")
            if any(term in text for term in conflict_terms):
                return False
        if event_profile.get("boat_event"):
            conflict_terms = ("人员亲水", "滩涂游玩", "人员涉水", "疑似人员")
            if any(term in text for term in conflict_terms):
                return False
        return True

    @staticmethod
    def event_profile(workflow_insight: dict[str, Any]) -> dict[str, Any]:
        event_name = str(workflow_insight.get("event_name") or "")
        event_code = str(workflow_insight.get("event_code") or "")
        event_text = f"{event_name} {event_code}"
        person_event = any(token in event_text for token in ("人员", "亲水", "涉水", "滩涂", "入侵", "PERSON_"))
        boat_event = any(token in event_text for token in ("船", "捕鱼", "电鱼", "偷捕", "BOAT_", "ILLEGAL_FISHING"))
        return {
            "event_name": event_name,
            "event_code": event_code,
            "person_event": person_event,
            "boat_event": boat_event,
        }

    def clean_report_text(self, text: str) -> str:
        text = str(text or "").strip()
        text = text.replace("✅", "")
        text = text.replace("▶", "")
        text = text.replace("—\n", "")
        text = self.strip_model_thinking(text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text or "未生成详细分析内容。"

    def compact(self, value: str, limit: int) -> str:
        text = self.clean_model_output_field(value)
        text = re.sub(r"\s+", " ", str(text or "")).strip()
        if len(text) <= limit:
            return text or "—"
        return self.sentence_safe_limit(text, limit)

    @staticmethod
    def sentence_safe_limit(value: Any, limit: int) -> str:
        text = re.sub(r"\s+", " ", str(value or "")).strip()
        if not text:
            return "—"
        if len(text) <= limit:
            return text
        truncated = text[: max(1, limit)].rstrip()
        cut_positions = [truncated.rfind(mark) for mark in ("。", "；", ";", "，", ",")]
        cut = max(cut_positions)
        if cut >= max(20, int(limit * 0.55)):
            sentence = truncated[: cut + 1].rstrip()
            if sentence.endswith(("，", ",", "；", ";", "：", ":")):
                sentence = f"{sentence.rstrip('，,；;：:')}。"
            return sentence
        return f"{truncated.rstrip('，,；;：:。')}。"

    def clean_model_output_field(self, value: Any, *, allow_empty: bool = False) -> str:
        text = str(value or "").strip()
        text = self.strip_model_thinking(text)
        text = re.sub(r"^\s*(结论|处置结论|最终结论)\s*[:：]\s*", "", text)
        text = re.sub(r"(?<![A-Za-z])medium\s*或\s*high(?![A-Za-z])", "中风险或高风险", text, flags=re.IGNORECASE)
        text = re.sub(r"(?<![A-Za-z])low\s*或\s*medium(?![A-Za-z])", "低风险或中风险", text, flags=re.IGNORECASE)
        text = re.sub(r"(?<![A-Za-z])high(?![A-Za-z])", "高风险", text, flags=re.IGNORECASE)
        text = re.sub(r"(?<![A-Za-z])medium(?![A-Za-z])", "中风险", text, flags=re.IGNORECASE)
        text = re.sub(r"(?<![A-Za-z])low(?![A-Za-z])", "低风险", text, flags=re.IGNORECASE)
        text = re.sub(
            r"(?:并)?将(?:补充运行状态和)?知识库依据(?:\[\d+\])?(?:写入|记录到|写入到)(?:本次)?(?:事件)?报告",
            "同步记录现场复核和人员撤离结果",
            text,
        )
        text = re.sub(
            r"(?:模型|系统)(?:将|会)依据知识库(?:生成|完善)(?:本次)?报告",
            "结合现场证据和运行状态完成风险研判",
            text,
        )
        text = re.sub(r"\s+", " ", text).strip()
        text = self.normalize_punctuation(text)
        if self.looks_like_model_thinking(text):
            text = ""
        return text if text else ("" if allow_empty else "—")

    @staticmethod
    def normalize_punctuation(value: Any) -> str:
        text = str(value or "")
        if not text:
            return ""
        text = re.sub(r"，{2,}", "，", text)
        text = re.sub(r"。{2,}", "。", text)
        text = re.sub(r"；{2,}", "；", text)
        text = re.sub(r"、{2,}", "、", text)
        text = re.sub(r"([。！？])\s*[，,；;、]+", r"\1", text)
        text = re.sub(r"([，,；;、])\s*([。！？])", r"\2", text)
        text = re.sub(r"([；;])\s*[，,、]+", "；", text)
        text = re.sub(r"([，,、])\s*[；;]+", "；", text)
        text = re.sub(r"：\s*[：:]+", "：", text)
        return text.strip()

    def valid_report_text(self, value: Any) -> str:
        text = self.clean_model_output_field(value, allow_empty=True)
        if text in {"", "—", "-", "无", "暂无", "未提供"}:
            return ""
        return text

    @staticmethod
    def strip_model_thinking(text: str) -> str:
        value = str(text or "").strip()
        if not value:
            return ""
        value = re.sub(r"<think>.*?</think>", "", value, flags=re.IGNORECASE | re.DOTALL).strip()
        thinking_patterns = [
            r"(?is)^\s*(?:结论|处置结论|最终结论)?\s*[:：]?\s*thinking\s+process\s*[:：].*$",
            r"(?is)^\s*(?:结论|处置结论|最终结论)?\s*[:：]?\s*reasoning\s*[:：].*$",
            r"(?is)^\s*(?:结论|处置结论|最终结论)?\s*[:：]?\s*analysis\s+process\s*[:：].*$",
        ]
        for pattern in thinking_patterns:
            if re.search(pattern, value):
                return ""
        value = re.sub(r"(?is)\bthinking\s+process\s*[:：].*$", "", value).strip()
        value = re.sub(r"(?is)\breasoning\s*[:：].*$", "", value).strip()
        return value

    @staticmethod
    def safe_filename_part(value: str, limit: int = 50) -> str:
        """清洗事件名，用作报告文件名前缀（去除文件系统非法字符）。"""
        text = re.sub(r'[\\/:*?"<>|\s]+', "_", str(value or "").strip())
        text = text.strip("_. ")
        return text[:limit] or "事件"

    def format_datetime(self, value: Optional[dt.datetime]) -> str:
        if not value:
            return "—"
        return self.to_local_datetime(value).strftime("%Y-%m-%d %H:%M:%S")

    @staticmethod
    def format_chinese_date(value: Any) -> str:
        if isinstance(value, dt.datetime):
            value = value.date()
        if isinstance(value, dt.date):
            return f"{value.year}年{value.month}月{value.day}日"
        text = str(value or "").strip()
        match = re.match(r"^(\d{4})[-/年](\d{1,2})[-/月](\d{1,2})", text)
        if match:
            year, month, day = (int(part) for part in match.groups())
            return f"{year}年{month}月{day}日"
        return text

    def to_local_datetime(self, value: dt.datetime) -> dt.datetime:
        """Convert DB/application timestamps to local display time.

        The backend container currently runs in UTC, so naïve application
        datetimes in event tables are UTC values. Reports are user-facing and
        should display Beijing time.
        """
        if value.tzinfo is None:
            current_utc_naive = dt.datetime.utcnow()
            if value <= current_utc_naive + dt.timedelta(hours=1):
                value = value.replace(tzinfo=dt.timezone.utc)
            else:
                value = value.replace(tzinfo=LOCAL_TIMEZONE)
        return value.astimezone(LOCAL_TIMEZONE)

    def workflow_insight(
        self,
        workflow_payload: dict[str, Any],
        visual: dict[str, Any],
        selected_text: str,
    ) -> dict[str, Any]:
        screening = visual.get("screening") if isinstance(visual.get("screening"), dict) else {}
        possible_flag = lambda key: (  # noqa: E731
            str(screening.get(key)) in {"1", "True", "true"}
            or screening.get(key) == 1
            or screening.get(key) is True
        )
        result = {
            "qwen_summary": screening.get("qwen_summary") or screening.get("summary"),
            "qwen_risk_level": screening.get("qwen_risk_level"),
            "flood_detected": screening.get("flood_detected"),
            "person_present": screening.get("person_present"),
            "boat_present": screening.get("boat_present"),
            "possible_person": screening.get("possible_person"),
            "possible_boat": screening.get("possible_boat"),
            "suspected": bool(possible_flag("possible_person") or possible_flag("possible_boat")),
            "raw_excerpt": self.extract_useful_excerpt(selected_text),
        }
        classify = self.find_node_inference(workflow_payload, "action_classify")
        main_class = classify.get("main_class")
        confidence = self.numeric_value(classify.get("confidence"))
        if confidence is None:
            confidence = self.best_key_frame_confidence(classify.get("key_frames"))
        result.update({
            "specialized_class": main_class,
            "specialized_class_label": self.class_label(main_class),
            "specialized_confidence": confidence,
            "sampled_frames": classify.get("sampled_frames"),
            "classification_report": classify.get("report"),
        })
        reasoning = self.find_node_inference(workflow_payload, "action_reasoning")
        knowledge_sources = self.find_in_value(reasoning, "knowledge_sources")
        sentence_citations = self.find_in_value(reasoning, "sentence_citations")
        result.update({
            "qwen4b_detailed_scene_analysis": self.find_in_value(reasoning, "detailed_scene_analysis"),
            "qwen4b_risk_reasoning": self.find_in_value(reasoning, "risk_reasoning"),
            "qwen4b_impact_assessment": self.find_in_value(reasoning, "impact_assessment"),
            "qwen4b_response_plan": self.find_in_value(reasoning, "response_plan"),
            "qwen4b_monitoring_suggestions": self.find_in_value(reasoning, "monitoring_suggestions"),
            "knowledge_sources": knowledge_sources if isinstance(knowledge_sources, list) else [],
            "sentence_citations": sentence_citations if isinstance(sentence_citations, list) else [],
            "knowledge_sources_summary": self.format_knowledge_sources(knowledge_sources, sentence_citations),
            "qwen4b_conclusion": (
                self.find_in_value(reasoning, "impact_assessment")
                or self.find_in_value(reasoning, "monitoring_suggestions")
            ),
        })
        # Keep citations emitted by the cloud reviewer in the same audit trail
        # as the edge citations. The source list is de-duplicated so the report
        # basis remains readable when both nodes cite the same clauses.
        cloud_reasoning = self.find_node_inference(workflow_payload, "action_report")
        cloud_sources = self.find_in_value(cloud_reasoning, "knowledge_sources")
        cloud_citations = self.find_in_value(cloud_reasoning, "sentence_citations")
        if isinstance(cloud_sources, list):
            merged_sources = []
            seen_sources = set()
            for source in [*(result.get("knowledge_sources") or []), *cloud_sources]:
                if not isinstance(source, dict):
                    continue
                source_key = str(
                    source.get("evidence_id")
                    or source.get("clause_id")
                    or source.get("chunk_id")
                    or ""
                ).strip()
                if not source_key or source_key in seen_sources:
                    continue
                seen_sources.add(source_key)
                merged_sources.append(source)
            result["knowledge_sources"] = merged_sources
        if isinstance(cloud_citations, list):
            result["sentence_citations"] = [
                *(result.get("sentence_citations") or []),
                *cloud_citations,
            ]
        if isinstance(cloud_sources, list) or isinstance(cloud_citations, list):
            result["knowledge_sources_summary"] = self.format_knowledge_sources(
                result.get("knowledge_sources") or [],
                result.get("sentence_citations"),
            )
        return result

    def report_knowledge_citations(
        self,
        workflow_payload: dict[str, Any],
        selected: dict[str, Any],
        workflow_insight: dict[str, Any],
    ) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
        """Return only citations used by the selected final report node.

        Edge citations remain in ``*_all`` for audit, but they must not make
        the final DOCX knowledge section claim support that the selected cloud
        report did not actually use.
        """
        all_sources = [
            item for item in workflow_insight.get("knowledge_sources") or []
            if isinstance(item, dict)
        ]
        node_id = selected.get("node_id") or "action_reasoning"
        node = self.find_node_inference(workflow_payload, str(node_id))
        node_citations = self.find_in_value(node, "sentence_citations")
        raw_citations = [item for item in node_citations if isinstance(item, dict)] if isinstance(node_citations, list) else []

        # A final node sometimes puts [K...] directly in a report field but
        # omits the corresponding sentence_citations entry. Accept only those
        # markers from the selected node, and resolve them against the complete
        # retrieved source list; never use edge-only citations here.
        text_parts: list[str] = []
        rendered_fields: dict[str, list[str]] = {}
        for key in (
            "detailed_scene_analysis",
            "risk_reasoning",
            "impact_assessment",
            "response_plan",
            "monitoring_suggestions",
            "key_observation",
            "conclusion",
        ):
            value = self.find_in_selected(selected, key)
            if value not in (None, "", []):
                value_text = str(value)
                text_parts.append(value_text)
                rendered_fields.setdefault(key, []).append(value_text)
        # ``selected.text`` may be a transport summary containing nested
        # template_data/handling_summary content that is not rendered as a
        # report paragraph. Only use it when no formal report field exists.
        if not text_parts and selected.get("text"):
            text_parts.append(str(selected.get("text")))
        selected_text_ids = {
            match.group(0)
            for text in text_parts
            for match in re.finditer(r"\bK\d+\b", text)
        }
        # A structured sentence_citations entry is not enough by itself. The
        # final report must visibly indicate the knowledge basis in the same
        # rendered field. The wording may paraphrase the source clause; it
        # does not need to quote the knowledge-base text verbatim.
        citations = []
        selected_ids: set[str] = set(selected_text_ids)
        provenance_terms = ("知识库", "知识依据", "条款", "依据", "规定", "规范")
        for citation in raw_citations:
            evidence_ids = {
                str(evidence_id).strip()
                for evidence_id in citation.get("evidence_ids") or []
                if str(evidence_id).strip()
            }
            sentence_ids = {
                match.group(0)
                for match in re.finditer(r"\bK\d+\b", str(citation.get("sentence") or ""))
            }
            citation_sentence = str(citation.get("sentence") or "").strip()
            field_name = str(citation.get("field") or "").strip()
            candidate_texts = rendered_fields.get(field_name) or text_parts
            normalized_sentence = re.sub(r"\s+", "", citation_sentence)
            sentence_visible = bool(normalized_sentence) and any(
                normalized_sentence in re.sub(r"\s+", "", text)
                for text in candidate_texts
            )
            # With an explicit [K...] marker, the marker must also be present
            # in the rendered text. Without it, accept a paraphrased citation
            # only when the citation sentence explicitly identifies the source
            # as a knowledge-base rule/requirement.
            if sentence_ids:
                visible_ids = evidence_ids & sentence_ids & selected_text_ids if sentence_visible else set()
            elif sentence_visible and any(term in citation_sentence for term in provenance_terms):
                visible_ids = evidence_ids
            else:
                visible_ids = set()
            if not visible_ids:
                continue
            citations.append(citation)
            selected_ids.update(visible_ids)
        sources = []
        for source in all_sources:
            identifiers = {
                str(source.get("evidence_id") or "").strip(),
                str(source.get("clause_id") or "").strip(),
            }
            chunk_id = str(source.get("chunk_id") or "").strip()
            if chunk_id:
                identifiers.add(f"K{chunk_id}")
            if identifiers & selected_ids:
                sources.append(source)
        return sources, citations

    @staticmethod
    def format_knowledge_sources(sources: Any, sentence_citations: Any = None) -> str:
        if not isinstance(sources, list):
            return ""
        citation_ids: set[str] = set()
        if isinstance(sentence_citations, list):
            for citation in sentence_citations:
                if not isinstance(citation, dict):
                    continue
                for evidence_id in citation.get("evidence_ids") or []:
                    value = str(evidence_id or "").strip()
                    if value:
                        citation_ids.add(value)

        # The retrieval list contains candidates; the report's knowledge basis
        # should show only clauses the model actually cited in its正文. Keep
        # the full list in workflow context for reasoning and audit purposes.
        if isinstance(sentence_citations, list):
            if not citation_ids:
                return "模型未在正文中明确引用知识库条款"
            selected_sources = []
            for item in sources:
                if not isinstance(item, dict):
                    continue
                identifiers = {
                    str(item.get("evidence_id") or "").strip(),
                    str(item.get("clause_id") or "").strip(),
                }
                chunk_id = str(item.get("chunk_id") or "").strip()
                if chunk_id:
                    identifiers.add(f"K{chunk_id}")
                if identifiers & citation_ids:
                    selected_sources.append(item)
            sources = selected_sources

        lines = []
        seen = set()
        source_count = 0
        for item in sources:
            if not isinstance(item, dict):
                continue
            title = item.get("document_title") or item.get("filename")
            if not title:
                continue
            section = item.get("section_path") or item.get("section_title") or ""
            clause = item.get("clause_id") or item.get("evidence_id") or ""
            key = (str(title).strip(), str(section).strip(), str(clause).strip())
            if key in seen:
                continue
            seen.add(key)
            source_count += 1
            title_text = str(title).strip()
            section_text = str(section).strip()
            # The root section is often identical to the document title.
            # Show only the document name there; for nested paths keep the
            # readable child section without repeating the title.
            if section_text == title_text:
                section_text = ""
            elif section_text.startswith(f"{title_text} > "):
                section_text = section_text[len(title_text) + 3:]
            line = f"[{source_count}] 《{title_text}》："
            if section_text:
                line += f"章节 {section_text}"
            if clause:
                line += f"{'，' if section_text else ''}条款 {str(clause).strip()}"
            lines.append(line)
            if source_count >= 5:
                break
        return "\n".join(lines)

    def find_node_inference(self, workflow_payload: dict[str, Any], node_id: str) -> dict[str, Any]:
        execution = workflow_payload.get("execution_result") if isinstance(workflow_payload, dict) else {}
        for row in (execution or {}).get("node_results") or []:
            if not isinstance(row, dict) or row.get("node_id") != node_id:
                continue
            output = row.get("output") if isinstance(row.get("output"), dict) else {}
            inference = output.get("inference_result")
            return inference if isinstance(inference, dict) else output
        return {}

    def numeric_value(self, value: Any) -> Optional[float]:
        try:
            if value in (None, "") or isinstance(value, (dict, list)):
                return None
            return float(value)
        except (TypeError, ValueError):
            return None

    def best_key_frame_confidence(self, key_frames: Any) -> Optional[float]:
        values = []
        if isinstance(key_frames, list):
            for frame in key_frames:
                if isinstance(frame, dict):
                    current = self.numeric_value(frame.get("confidence"))
                    if current is not None:
                        values.append(current)
        return max(values) if values else None

    def class_label(self, value: Any) -> str:
        labels = {
            "flood": "洪水",
            "landslide": "滑坡",
            "mudslide": "泥石流",
            "earthquake": "地震",
            "person": "人员",
            "boat": "船只",
        }
        text = str(value or "").strip()
        return labels.get(text.lower(), text)

    def extract_useful_excerpt(self, text: str) -> str:
        cleaned = self.clean_report_text(text)
        if "三、分析流程规划" in cleaned and not cleaned.rstrip().endswith("。"):
            cleaned = cleaned.split("三、分析流程规划", 1)[0].strip()
        cleaned = re.sub(r"【[^】]+】", "", cleaned).strip()
        return self.compact(cleaned, 520)

    def handling_summary(
        self,
        *,
        instance: SafetyEventInstance,
        event: EventLibrary,
        visual: dict[str, Any],
        selected: dict[str, Any],
        selected_text: str,
        workflow_insight: dict[str, Any],
        image_items: list[dict[str, Any]],
        video_items: list[dict[str, Any]],
    ) -> str:
        suspect_note = self.screening_suspect_note(visual)
        detailed = self.selected_detailed_summary(selected)
        if detailed:
            if suspect_note and suspect_note not in detailed:
                return f"一、初筛复核：{suspect_note}\n{detailed}"
            return detailed
        detailed = self.workflow_detailed_summary(workflow_insight)
        if detailed:
            if suspect_note and suspect_note not in detailed:
                return f"一、初筛复核：{suspect_note}\n{detailed}"
            return detailed

        event_name = getattr(event, "event_name", None) or instance.summary or "安全事件"
        qwen_summary = workflow_insight.get("qwen_summary") or "初筛未提供明确文字摘要"
        class_label = workflow_insight.get("specialized_class_label") or "未获得专有模型类别"
        confidence = workflow_insight.get("specialized_confidence")
        confidence_text = f"{confidence * 100:.1f}%" if confidence is not None else "—"
        sampled_frames = workflow_insight.get("sampled_frames") or len(image_items) or "—"
        cloud_text = ""
        if selected.get("cloud_error"):
            cloud_text = "云端增强暂不可用，本报告采用本地 4B 场景理解与专有模型结果整理生成。"
        else:
            cloud_text = f"报告来源为{selected.get('source_label', '智能分析模型')}。"

        risk_text = RISK_NAMES.get(str(instance.max_risk_level or instance.risk_level or "").upper(), "待确认")
        lines = [
            f"一、事件复核：系统触发{event_name}，当前风险等级为{risk_text}。{cloud_text}",
            f"二、现场证据：{suspect_note or f'摄像头初筛摘要为“{qwen_summary}”。'}本次关联事件证据视频{len(video_items)}段、关键帧/检测图像{len(image_items)}张。",
            f"三、模型研判：专有模型复核类别为{class_label}，置信度{confidence_text}，采样帧数{sampled_frames}。边缘侧4B已结合视频证据、事件类型和上下文完成复核。",
            "四、处置建议：保持摄像头连续取证，联动相关传感器和现场巡查记录；如风险指标持续升高或现场出现人员、设施受威胁情况，应升级告警并启动现场处置。",
        ]
        return "\n".join(lines)

    def selected_detailed_summary(self, selected: dict[str, Any]) -> str:
        fields = [
            ("一、现场场景", self.final_report_field(selected, "detailed_scene_analysis", "")),
            ("二、风险研判", self.final_report_field(selected, "risk_reasoning", "")),
            ("三、影响评估", self.final_report_field(selected, "impact_assessment", "")),
            ("四、处置建议", self.final_report_field(selected, "response_plan", "")),
            ("五、持续监测", self.final_report_field(selected, "monitoring_suggestions", "")),
        ]
        lines = [
            f"{label}：{text}"
            for label, text in fields
            if text and text != "—" and not self.looks_like_model_thinking(text)
        ]
        return "\n".join(lines) if len(lines) >= 2 else ""

    @staticmethod
    def workflow_detailed_summary(workflow_insight: dict[str, Any]) -> str:
        fields = [
            ("一、现场场景", workflow_insight.get("qwen4b_detailed_scene_analysis")),
            ("二、风险研判", workflow_insight.get("qwen4b_risk_reasoning")),
            ("三、影响评估", workflow_insight.get("qwen4b_impact_assessment")),
            ("四、处置建议", workflow_insight.get("qwen4b_response_plan")),
            ("五、持续监测", workflow_insight.get("qwen4b_monitoring_suggestions")),
        ]
        lines = [f"{label}：{str(text).strip()}" for label, text in fields if str(text or "").strip()]
        return "\n".join(lines) if len(lines) >= 2 else ""

    @staticmethod
    def looks_like_model_thinking(text: str) -> bool:
        value = str(text or "").strip()
        lowered = value.lower()
        return (
            value.startswith(("好的，我现在", "首先，我", "<think>"))
            or lowered.startswith(("thinking process", "reasoning:", "analysis process", "the user wants", "we need to", "i need to"))
            or "thinking process:" in lowered[:80]
            or "analyze the request" in lowered[:160]
            or "the input data includes" in lowered[:220]
            or "</think>" in value
            or "需要处理用户" in value[:120]
            or "生成一个关于" in value[:160]
        )

    def looks_truncated(self, text: str) -> bool:
        stripped = str(text or "").strip()
        if not stripped:
            return True
        if stripped.endswith(("专", "“", "：", ":", "，", "、", "-", "—")):
            return True
        return len(stripped) > 80 and not re.search(r"[。！？.!?]$", stripped)


dam_event_report_service = DamEventReportService()
